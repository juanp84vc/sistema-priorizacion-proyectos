#!/usr/bin/env python3
"""
Script to update the Guía Operativa with CONFIS methodology integration
Version 2: More surgical approach with better paragraph tracking
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from copy import deepcopy

def inspect_document(doc, start=0, end=None):
    """Print document structure for inspection"""
    if end is None:
        end = len(doc.paragraphs)

    for i in range(start, min(end, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text[:80].replace('\n', ' ')
        print(f"{i:3d}: {text}")

def find_paragraphs_by_text(doc, search_text, partial=True, case_sensitive=False):
    """Find all paragraphs containing specific text"""
    results = []
    search = search_text if case_sensitive else search_text.lower()

    for i, para in enumerate(doc.paragraphs):
        text = para.text if case_sensitive else para.text.lower()
        if partial:
            if search in text:
                results.append((i, para))
        else:
            if text == search:
                results.append((i, para))
    return results

def clear_paragraph_content(para):
    """Clear all runs from a paragraph"""
    for run in para.runs:
        r = run._element
        r.getparent().remove(r)

def set_paragraph_text(para, new_text, bold=False, italic=False):
    """Set paragraph text with optional formatting"""
    clear_paragraph_content(para)
    run = para.add_run(new_text)
    run.bold = bold
    run.italic = italic

def insert_paragraph_after(doc, para_index, text, preserve_style=True):
    """Insert a new paragraph after the given index"""
    old_para = doc.paragraphs[para_index]
    new_para = old_para._element
    p = deepcopy(new_para)
    new_para.addnext(p)

    # Set the text
    new_p_obj = doc.paragraphs[para_index + 1]
    set_paragraph_text(new_p_obj, text)
    return para_index + 1

def update_guia_v2(doc_path):
    """Main update function - version 2"""

    doc = Document(doc_path)
    print(f"Document loaded. Total paragraphs: {len(doc.paragraphs)}\n")

    # ========================================
    # UPDATE 1: Header/Version (paragraph 8)
    # ========================================
    print("[1] UPDATING HEADER/VERSION")
    print("-" * 80)

    results = find_paragraphs_by_text(doc, "Incluye ajustes metodológicos", partial=True)
    if results:
        idx, para = results[0]
        old = para.text
        new = "Incluye ajustes metodológicos: SROI logarítmico, rúbricas objetivas, alertas contextuales, metodología CONFIS"
        set_paragraph_text(para, new)
        print(f"✓ Para {idx}: Updated subtitle")
        print(f"  Old: {old}")
        print(f"  New: {new}\n")
    else:
        print("✗ Could not find subtitle paragraph\n")

    # ========================================
    # UPDATE 2: Section 3.4 Alcance Territorial
    # ========================================
    print("[2] UPDATING SECTION 3.4 - ALCANCE TERRITORIAL")
    print("-" * 80)

    # Find the section
    alcance_results = find_paragraphs_by_text(doc, "Alcance Territorial", partial=True)
    if alcance_results:
        alcance_idx, _ = alcance_results[0]
        print(f"Found 'Alcance Territorial' at paragraph {alcance_idx}")

        # Find lines to update (look for "Base:" and "Bonus:" patterns)
        base_results = find_paragraphs_by_text(doc, "Base: 10 pts por municipio", partial=True)

        if base_results:
            base_idx, base_para = base_results[0]
            print(f"Found 'Base: 10 pts por municipio' at paragraph {base_idx}")

            # Replace Base line
            set_paragraph_text(base_para, "Base: 10 pts por municipio (máx. 30 pts)")
            print(f"✓ Para {base_idx}: Updated Base line")

            # Look for bonus lines after base
            pdet_results = find_paragraphs_by_text(doc, "Bonus PDET", partial=True)
            if pdet_results:
                pdet_idx, pdet_para = pdet_results[0]
                old_text = pdet_para.text
                set_paragraph_text(pdet_para, "Bonus PDET: +15 pts si incluye municipios PDET")
                print(f"✓ Para {pdet_idx}: Updated PDET bonus")
                print(f"  Old: {old_text}")

            # Look for the "Máximo" or next bonus after PDET
            if pdet_results:
                pdet_idx, _ = pdet_results[0]
                # Check next few paragraphs
                for i in range(pdet_idx + 1, min(pdet_idx + 5, len(doc.paragraphs))):
                    para = doc.paragraphs[i]
                    if "Bonus" in para.text or "máx" in para.text.lower():
                        if "multi" not in para.text.lower():
                            # This is where we add multi-departamental
                            set_paragraph_text(para, "Bonus multi-departamental: +15 pts")
                            print(f"✓ Para {i}: Updated/added multi-departamental bonus")

                            # Check if there's a corredor bonus
                            if i + 1 < len(doc.paragraphs):
                                next_para = doc.paragraphs[i + 1]
                                if "máx" not in next_para.text.lower() and "Máximo" not in next_para.text:
                                    set_paragraph_text(next_para, "Bonus corredor: +10 pts")
                                    print(f"✓ Para {i + 1}: Updated/added corredor bonus")

                                    # Add máximo line
                                    if i + 2 < len(doc.paragraphs):
                                        max_para = doc.paragraphs[i + 2]
                                        set_paragraph_text(max_para, "Máximo posible: 30 + 30 + 15 + 15 + 10 = 100 pts")
                                        print(f"✓ Para {i + 2}: Updated máximo line")
                            break
    print()

    # ========================================
    # UPDATE 3: Section 4 - Criterio 3 Title and Description
    # ========================================
    print("[3] UPDATING SECTION 4 - CRITERIO 3")
    print("-" * 80)

    criterio3_results = find_paragraphs_by_text(doc, "Criterio 3", partial=True)
    if criterio3_results:
        for idx, para in criterio3_results:
            if "Probabilidad de Aprobación" in para.text:
                old = para.text
                new = "4. Criterio 3: Probabilidad de Aprobación CONFIS (20%)"
                set_paragraph_text(para, new)
                print(f"✓ Para {idx}: Updated title")
                print(f"  Old: {old}")
                print(f"  New: {new}\n")

                # Update the description paragraph
                if idx + 1 < len(doc.paragraphs):
                    desc_para = doc.paragraphs[idx + 1]
                    # Only update if it looks like description
                    if len(desc_para.text) > 20:
                        old_desc = desc_para.text
                        new_desc = "Evalúa la probabilidad de aprobación usando la metodología CONFIS (Anexo 2). Integra grupo de priorización, puntaje territorial y puntaje sectorial."
                        set_paragraph_text(desc_para, new_desc)
                        print(f"✓ Para {idx + 1}: Updated description")
                        print(f"  New: {new_desc}\n")
                break
    else:
        print("✗ Could not find 'Criterio 3' section\n")

    # ========================================
    # UPDATE 4: Section 4.1 - Fórmula CONFIS
    # ========================================
    print("[4] UPDATING SECTION 4.1 - FÓRMULA CONFIS")
    print("-" * 80)

    factores_results = find_paragraphs_by_text(doc, "Factores que considera", partial=True)
    if factores_results:
        idx, para = factores_results[0]
        old = para.text
        set_paragraph_text(para, "4.1 Fórmula CONFIS")
        print(f"✓ Para {idx}: Updated subsection title")
        print(f"  Old: {old}")
        print(f"  New: 4.1 Fórmula CONFIS\n")

        # Replace formula content in next paragraphs
        if idx + 1 < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx + 1], "Score = GrupoPrioridad × 20% + ScoreCONFIS_Normalizado × 80%")
            print(f"✓ Para {idx + 1}: Added formula line 1")

        if idx + 2 < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[idx + 2], "ScoreCONFIS = (PuntajeTerritorial + PuntajeSectorial) / 20 × 100")
            print(f"✓ Para {idx + 2}: Added formula line 2")
    else:
        print("✗ Could not find 'Factores que considera' section\n")

    # ========================================
    # UPDATE 5: CONFIS Groups explanation
    # ========================================
    print("\n[5] UPDATING CONFIS GROUPS AND ELEGIBILITY GATE")
    print("-" * 80)

    # Look for bonus text that needs to be replaced with groups info
    bonus_results = find_paragraphs_by_text(doc, "Bonus PDET", partial=True)
    if bonus_results:
        for idx, para in bonus_results:
            # Skip if already updated (contains "si incluye")
            if "si incluye municipios PDET" in para.text:
                continue

            # This might be an old bonus that needs groups info
            # Check context - look for paragraph around 60-65
            if 60 <= idx <= 75:
                # Look ahead for a paragraph that could hold groups info
                for i in range(idx, min(idx + 8, len(doc.paragraphs))):
                    p = doc.paragraphs[i]
                    if "Bonus" in p.text and i > idx:
                        # Found a place to put groups info
                        groups_text = "8 Grupos de priorización: Grupo 1-2 (PATR-PDET), 3-4 (PDET), 5-6 (ZOMAC), 7-8 (Amazonía). Grupos impares = contribuyente paga estructuración."
                        set_paragraph_text(p, groups_text)
                        print(f"✓ Para {i}: Added CONFIS groups explanation")

                        # Add gate de elegibilidad after
                        if i + 1 < len(doc.paragraphs):
                            gate_para = doc.paragraphs[i + 1]
                            gate_text = "Gate de elegibilidad: Solo proyectos en municipios PDET/ZOMAC/Amazonía son elegibles. Proyectos fuera reciben score 0 y nivel NO ELEGIBLE."
                            set_paragraph_text(gate_para, gate_text)
                            print(f"✓ Para {i + 1}: Added gate de elegibilidad\n")
                        break

    # ========================================
    # Save document
    # ========================================
    print("\n" + "=" * 80)
    print("SAVING DOCUMENT")
    print("=" * 80)
    doc.save(doc_path)
    print(f"✓ Document saved: {doc_path}\n")

if __name__ == "__main__":
    doc_path = "/sessions/keen-stoic-bell/mnt/sistema-priorizacion-proyectos/Guia_Operativa_Evaluadores_ENLAZA_GEB.docx"

    try:
        update_guia_v2(doc_path)
        print("✓ ALL UPDATES COMPLETED SUCCESSFULLY!")
    except Exception as e:
        print(f"✗ ERROR: {e}")
        import traceback
        traceback.print_exc()
