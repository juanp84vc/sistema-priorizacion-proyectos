#!/usr/bin/env python3
"""
Final surgical update script for CONFIS methodology integration
Based on actual document structure inspection
"""

from docx import Document

def set_paragraph_text(para, new_text):
    """Clear and set paragraph text"""
    # Remove all runs
    for run in para.runs:
        r = run._element
        r.getparent().remove(r)
    # Add new text
    para.add_run(new_text)

def update_document(doc_path):
    """Apply all CONFIS updates"""

    doc = Document(doc_path)
    print(f"Document loaded: {len(doc.paragraphs)} paragraphs\n")

    # ========================================
    # UPDATE 1: Header subtitle (para 8)
    # ========================================
    print("[1] UPDATING HEADER SUBTITLE (para 8)")
    print("-" * 80)

    para_8 = doc.paragraphs[8]
    old_text = para_8.text
    new_text = "Incluye ajustes metodológicos: SROI logarítmico, rúbricas objetivas, alertas contextuales, metodología CONFIS"
    set_paragraph_text(para_8, new_text)
    print(f"✓ Updated paragraph 8")
    print(f"  Old: {old_text}")
    print(f"  New: {new_text}\n")

    # ========================================
    # UPDATE 2: Section 3.4 Alcance Territorial
    # ========================================
    print("[2] UPDATING SECTION 3.4 - ALCANCE TERRITORIAL")
    print("-" * 80)

    # Para 57: Add the CONFIS puntaje line
    para_57 = doc.paragraphs[57]
    para_57_old = para_57.text
    # This might already be updated, check first
    if "CONFIS" not in para_57.text:
        set_paragraph_text(para_57, "Puntaje Territorial CONFIS: puntaje × 3 (máx. 30 pts)")
        print(f"✓ Para 57: Added CONFIS puntaje line")
    else:
        print(f"ℹ Para 57: Already contains CONFIS reference")

    # Para 59: Already has Base line (verified in inspection)
    para_59 = doc.paragraphs[59]
    if "máx. 30 pts)" not in para_59.text:
        set_paragraph_text(para_59, "Base: 10 pts por municipio (máx. 30 pts)")
        print(f"✓ Para 59: Updated Base line")
    else:
        print(f"ℹ Para 59: Already correct")

    # Para 60: Groups explanation (already added in previous run)
    para_60 = doc.paragraphs[60]
    if "8 Grupos de priorización" not in para_60.text:
        groups_text = "8 Grupos de priorización: Grupo 1-2 (PATR-PDET), 3-4 (PDET), 5-6 (ZOMAC), 7-8 (Amazonía). Grupos impares = contribuyente paga estructuración."
        set_paragraph_text(para_60, groups_text)
        print(f"✓ Para 60: Added CONFIS groups explanation")
    else:
        print(f"ℹ Para 60: Already contains groups info")

    # Para 62: Update with Bonus PDET (was "Bonus multi-departamental")
    para_62 = doc.paragraphs[62]
    if "Bonus PDET:" not in para_62.text:
        set_paragraph_text(para_62, "Bonus PDET: +15 pts si incluye municipios PDET")
        print(f"✓ Para 62: Updated to PDET bonus")
    else:
        print(f"ℹ Para 62: Already PDET bonus")

    # Insert bonus corredor and corredor lines if not present
    # Check if we need to add them between current para 62 and para 63
    para_63_text = doc.paragraphs[63].text
    if "Criterio 3" not in para_63_text:
        # We still have room, add the bonuses before criterio 3
        pass
    else:
        # Criterio 3 is at para 63, so we need to insert before it
        # Let's add the missing bonus lines in the right places

        # Insert "Bonus multi-departamental: +15 pts" before para 62
        if "Bonus multi-departamental" not in doc.paragraphs[61].text:
            # Create new paragraph with multi-departamental bonus
            # For now, update para 61 if empty
            para_61 = doc.paragraphs[61]
            if len(para_61.text.strip()) < 10:  # If nearly empty
                set_paragraph_text(para_61, "Bonus multi-departamental: +15 pts")
                print(f"✓ Para 61: Added multi-departamental bonus")

    # Add corredor bonus
    # Find a good place - check para 62 area
    # Let's insert after the multi-departamental bonus
    insert_after_62 = False
    for i in range(62, 68):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if "Bonus corredor" not in para.text and "Criterio 3" not in para.text and "Máximo" not in para.text:
                if "Bonus ZOMAC" in para.text or "Bonus multi" in para.text or len(para.text) < 30:
                    # Replace or update this paragraph
                    set_paragraph_text(para, "Bonus corredor: +10 pts")
                    print(f"✓ Para {i}: Added corredor bonus")
                    break

    # Add gate de elegibilidad
    for i in range(62, 72):
        if i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if "Gate de elegibilidad" not in para.text and len(para.text) < 50:
                if "Bonus" not in para.text and "Criterio" not in para.text:
                    set_paragraph_text(para, "Gate de elegibilidad: Solo proyectos en municipios PDET/ZOMAC/Amazonía son elegibles. Proyectos fuera reciben score 0 y nivel NO ELEGIBLE.")
                    print(f"✓ Para {i}: Added gate de elegibilidad")
                    break

    print()

    # ========================================
    # UPDATE 3: Section 4.1 - Fórmula CONFIS
    # ========================================
    print("[3] UPDATING SECTION 4.1 - FÓRMULA CONFIS")
    print("-" * 80)

    # Para 65: Already has "4.1 Fórmula CONFIS"
    para_65 = doc.paragraphs[65]
    if "Fórmula CONFIS" in para_65.text:
        print(f"✓ Para 65: Already has 'Fórmula CONFIS' title")
    else:
        set_paragraph_text(para_65, "4.1 Fórmula CONFIS")
        print(f"✓ Para 65: Updated to 'Fórmula CONFIS'")

    # Para 66 and 67: Add formula lines
    # Check what's there currently
    para_66 = doc.paragraphs[66]
    para_67 = doc.paragraphs[67]

    if len(para_66.text.strip()) == 0:
        set_paragraph_text(para_66, "Score = GrupoPrioridad × 20% + ScoreCONFIS_Normalizado × 80%")
        print(f"✓ Para 66: Added formula line 1")
    else:
        print(f"ℹ Para 66: {para_66.text[:50]}")

    if "Score" not in para_67.text:
        set_paragraph_text(para_67, "ScoreCONFIS = (PuntajeTerritorial + PuntajeSectorial) / 20 × 100")
        print(f"✓ Para 67: Added formula line 2")
    else:
        print(f"ℹ Para 67: {para_67.text[:50]}")

    print()

    # ========================================
    # Save document
    # ========================================
    print("=" * 80)
    print("SAVING DOCUMENT")
    print("=" * 80)
    doc.save(doc_path)
    print(f"✓ Document saved: {doc_path}\n")

    return doc

if __name__ == "__main__":
    doc_path = "Guia_Operativa_Evaluadores_ENLAZA_GEB.docx"

    try:
        doc = update_document(doc_path)
        print("✓ ALL UPDATES COMPLETED!")
        print("\nFinal structure verification:")
        print("-" * 80)
        for i in range(57, 70):
            if i < len(doc.paragraphs):
                text = doc.paragraphs[i].text[:75]
                print(f"Para {i:2d}: {text}")
    except Exception as e:
        print(f"✗ ERROR: {e}")
        import traceback
        traceback.print_exc()
