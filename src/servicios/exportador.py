"""
Módulo para exportar resultados de evaluación a diferentes formatos.
Soporta: Excel, CSV, Word y PDF.
"""
import io
from datetime import datetime
from typing import List, Dict, Any
import pandas as pd


class ExportadorResultados:
    """Clase para exportar resultados de evaluación a diferentes formatos."""

    def __init__(self, reporte: Dict[str, Any], resultados_detallados: List[Any] = None):
        """
        Inicializa el exportador con los datos del reporte.

        Args:
            reporte: Diccionario con el reporte completo de evaluación
            resultados_detallados: Lista opcional de resultados detallados por proyecto
        """
        self.reporte = reporte
        self.resultados_detallados = resultados_detallados or []
        self.fecha_generacion = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    @staticmethod
    def formatear_numero(numero: float, decimales: int = 2) -> str:
        """
        Formatea un número con punto para miles y coma para decimales.

        Args:
            numero: Número a formatear
            decimales: Cantidad de decimales a mostrar

        Returns:
            str: Número formateado (ej: 1.234.567,89)
        """
        if numero is None:
            return "0,00"

        # Formatear con decimales
        formato = f"{{:,.{decimales}f}}"
        numero_formateado = formato.format(numero)

        # Intercambiar punto y coma (de formato US a formato europeo/latinoamericano)
        numero_formateado = numero_formateado.replace(",", "TEMP")
        numero_formateado = numero_formateado.replace(".", ",")
        numero_formateado = numero_formateado.replace("TEMP", ".")

        return numero_formateado

    def exportar_csv(self) -> bytes:
        """
        Exporta el ranking a formato CSV.

        Returns:
            bytes: Contenido del archivo CSV
        """
        df = pd.DataFrame(self.reporte['ranking'])

        # Seleccionar columnas relevantes
        columnas = ['posicion', 'proyecto_id', 'proyecto_nombre', 'score', 'recomendacion']
        df_export = df[columnas].copy()

        # Formatear el score con punto para miles y coma para decimales
        df_export['score'] = df_export['score'].apply(lambda x: self.formatear_numero(x, 2))

        return df_export.to_csv(index=False, sep=';').encode('utf-8')

    def exportar_excel(self) -> bytes:
        """
        Exporta a Excel con múltiples hojas: Resumen, Ranking y Detalles.

        Returns:
            bytes: Contenido del archivo Excel
        """
        if not self.reporte or not self.reporte.get('ranking'):
            raise ValueError("No hay datos para exportar")

        output = io.BytesIO()

        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Hoja 1: Resumen Ejecutivo
            df_resumen = pd.DataFrame([
                ['Fecha de Generación', self.fecha_generacion],
                ['Total Proyectos Evaluados', str(self.reporte['total_proyectos'])],
                ['Score Máximo', self.formatear_numero(self.reporte['estadisticas']['score_maximo'])],
                ['Score Mínimo', self.formatear_numero(self.reporte['estadisticas']['score_minimo'])],
                ['Score Promedio', self.formatear_numero(self.reporte['estadisticas']['score_promedio'])],
                ['Desviación Estándar', self.formatear_numero(self.reporte['estadisticas']['desviacion_estandar'])],
                ['Proyectos Alta Prioridad', str(self.reporte['estadisticas']['proyectos_alta_prioridad'])],
                ['Proyectos Media Prioridad', str(self.reporte['estadisticas']['proyectos_media_prioridad'])],
                ['Proyectos Baja Prioridad', str(self.reporte['estadisticas']['proyectos_baja_prioridad'])],
            ], columns=['Métrica', 'Valor'])

            df_resumen.to_excel(writer, sheet_name='Resumen', index=False)

            # Hoja 2: Ranking
            df_ranking = pd.DataFrame(self.reporte['ranking']).copy()
            columnas_ranking = ['posicion', 'proyecto_id', 'proyecto_nombre', 'score', 'recomendacion']
            df_ranking_export = df_ranking[columnas_ranking].copy()

            # Formatear el score
            df_ranking_export['score'] = df_ranking_export['score'].apply(lambda x: self.formatear_numero(x, 2))

            df_ranking_export.to_excel(writer, sheet_name='Ranking', index=False)

            # Hoja 3: Detalles por Criterio
            if self.resultados_detallados:
                detalles_list = []

                for resultado in self.resultados_detallados:
                    for criterio, detalle in resultado.detalle_criterios.items():
                        detalles_list.append({
                            'Proyecto': resultado.proyecto_nombre,
                            'Criterio': criterio,
                            'Score Base': self.formatear_numero(detalle['score_base'], 2),
                            'Peso': self.formatear_numero(detalle['peso'] * 100, 1) + '%',
                            'Score Ponderado': self.formatear_numero(detalle['score_ponderado'], 2)
                        })

                df_detalles = pd.DataFrame(detalles_list)
                df_detalles.to_excel(writer, sheet_name='Detalles por Criterio', index=False)

            # Ajustar anchos de columna
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width

        return output.getvalue()

    def exportar_word(self) -> bytes:
        """
        Exporta a Word con formato profesional.

        Returns:
            bytes: Contenido del archivo Word
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("Se requiere python-docx. Instalar con: pip install python-docx")

        doc = Document()

        # Título
        titulo = doc.add_heading('Reporte de Evaluación de Proyectos', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información general
        doc.add_paragraph(f"Fecha de generación: {self.fecha_generacion}")
        doc.add_paragraph()

        # Sección: Resumen Ejecutivo
        doc.add_heading('Resumen Ejecutivo', 1)

        tabla_resumen = doc.add_table(rows=7, cols=2)
        tabla_resumen.style = 'Light Grid Accent 1'

        resumen_data = [
            ('Total Proyectos Evaluados', str(self.reporte['total_proyectos'])),
            ('Score Máximo', self.formatear_numero(self.reporte['estadisticas']['score_maximo'])),
            ('Score Promedio', self.formatear_numero(self.reporte['estadisticas']['score_promedio'])),
            ('Score Mínimo', self.formatear_numero(self.reporte['estadisticas']['score_minimo'])),
            ('Proyectos Alta Prioridad', str(self.reporte['estadisticas']['proyectos_alta_prioridad'])),
            ('Proyectos Media Prioridad', str(self.reporte['estadisticas']['proyectos_media_prioridad'])),
            ('Proyectos Baja Prioridad', str(self.reporte['estadisticas']['proyectos_baja_prioridad'])),
        ]

        for idx, (metrica, valor) in enumerate(resumen_data):
            tabla_resumen.rows[idx].cells[0].text = metrica
            tabla_resumen.rows[idx].cells[1].text = valor

        doc.add_paragraph()

        # Sección: Ranking de Proyectos
        doc.add_heading('Ranking de Proyectos', 1)

        # Crear tabla de ranking
        num_proyectos = len(self.reporte['ranking'])
        tabla_ranking = doc.add_table(rows=num_proyectos + 1, cols=4)
        tabla_ranking.style = 'Light Grid Accent 1'

        # Encabezados
        headers = ['Posición', 'Proyecto', 'Score', 'Recomendación']
        for idx, header in enumerate(headers):
            cell = tabla_ranking.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Datos
        for idx, proyecto in enumerate(self.reporte['ranking'], start=1):
            tabla_ranking.rows[idx].cells[0].text = str(proyecto['posicion'])
            tabla_ranking.rows[idx].cells[1].text = proyecto['proyecto_nombre']
            tabla_ranking.rows[idx].cells[2].text = self.formatear_numero(proyecto['score'])
            tabla_ranking.rows[idx].cells[3].text = proyecto['recomendacion']

        doc.add_paragraph()

        # Sección: Detalles por Proyecto
        if self.resultados_detallados:
            doc.add_heading('Detalles por Proyecto', 1)

            for resultado in self.resultados_detallados:
                doc.add_heading(f"{resultado.proyecto_nombre}", 2)
                doc.add_paragraph(f"Score Final: {self.formatear_numero(resultado.score_final)}")
                doc.add_paragraph(f"Recomendación: {resultado.recomendacion}")

                # Tabla de criterios
                tabla_criterios = doc.add_table(rows=len(resultado.detalle_criterios) + 1, cols=4)
                tabla_criterios.style = 'Light List Accent 1'

                # Encabezados
                tabla_criterios.rows[0].cells[0].text = 'Criterio'
                tabla_criterios.rows[0].cells[1].text = 'Score Base'
                tabla_criterios.rows[0].cells[2].text = 'Peso'
                tabla_criterios.rows[0].cells[3].text = 'Score Ponderado'

                for idx, (criterio, detalle) in enumerate(resultado.detalle_criterios.items(), start=1):
                    tabla_criterios.rows[idx].cells[0].text = criterio
                    tabla_criterios.rows[idx].cells[1].text = self.formatear_numero(detalle['score_base'])
                    tabla_criterios.rows[idx].cells[2].text = self.formatear_numero(detalle['peso'] * 100, 1) + '%'
                    tabla_criterios.rows[idx].cells[3].text = self.formatear_numero(detalle['score_ponderado'])

                # Observaciones
                if resultado.observaciones:
                    doc.add_paragraph("Observaciones:")
                    for obs in resultado.observaciones:
                        doc.add_paragraph(f"• {obs}", style='List Bullet')

                doc.add_paragraph()

        # Guardar en BytesIO
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def exportar_pdf(self) -> bytes:
        """
        Exporta a PDF con gráficos y formato profesional.

        Returns:
            bytes: Contenido del archivo PDF
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Table, TableStyle, Paragraph,
                Spacer, PageBreak, Image
            )
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError:
            raise ImportError("Se requiere reportlab. Instalar con: pip install reportlab")

        output = io.BytesIO()
        doc = SimpleDocTemplate(output, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12
        )

        # Título
        story.append(Paragraph("Reporte de Evaluación de Proyectos", title_style))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(f"Fecha de generación: {self.fecha_generacion}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))

        # Resumen Ejecutivo
        story.append(Paragraph("Resumen Ejecutivo", heading_style))

        resumen_data = [
            ['Métrica', 'Valor'],
            ['Total Proyectos Evaluados', str(self.reporte['total_proyectos'])],
            ['Score Máximo', self.formatear_numero(self.reporte['estadisticas']['score_maximo'])],
            ['Score Promedio', self.formatear_numero(self.reporte['estadisticas']['score_promedio'])],
            ['Score Mínimo', self.formatear_numero(self.reporte['estadisticas']['score_minimo'])],
            ['Desviación Estándar', self.formatear_numero(self.reporte['estadisticas']['desviacion_estandar'])],
            ['Proyectos Alta Prioridad', str(self.reporte['estadisticas']['proyectos_alta_prioridad'])],
            ['Proyectos Media Prioridad', str(self.reporte['estadisticas']['proyectos_media_prioridad'])],
            ['Proyectos Baja Prioridad', str(self.reporte['estadisticas']['proyectos_baja_prioridad'])],
        ]

        tabla_resumen = Table(resumen_data, colWidths=[3.5 * inch, 2 * inch])
        tabla_resumen.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
        ]))

        story.append(tabla_resumen)
        story.append(Spacer(1, 0.3 * inch))

        # Ranking de Proyectos
        story.append(Paragraph("Ranking de Proyectos", heading_style))

        ranking_data = [['Pos.', 'Proyecto', 'Score', 'Recomendación']]
        for proyecto in self.reporte['ranking']:
            ranking_data.append([
                str(proyecto['posicion']),
                proyecto['proyecto_nombre'][:40],
                self.formatear_numero(proyecto['score']),
                proyecto['recomendacion']
            ])

        tabla_ranking = Table(ranking_data, colWidths=[0.5 * inch, 2.5 * inch, 1 * inch, 2 * inch])
        tabla_ranking.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
        ]))

        story.append(tabla_ranking)
        story.append(Spacer(1, 0.3 * inch))

        # Detalles por Proyecto
        if self.resultados_detallados:
            story.append(PageBreak())
            story.append(Paragraph("Detalles por Proyecto", heading_style))

            for idx, resultado in enumerate(self.resultados_detallados):
                if idx > 0:
                    story.append(Spacer(1, 0.2 * inch))

                # Nombre del proyecto
                story.append(Paragraph(
                    f"<b>{resultado.proyecto_nombre}</b>",
                    styles['Heading3']
                ))
                story.append(Paragraph(
                    f"Score Final: {self.formatear_numero(resultado.score_final)} | {resultado.recomendacion}",
                    styles['Normal']
                ))
                story.append(Spacer(1, 0.1 * inch))

                # Tabla de criterios
                criterios_data = [['Criterio', 'Score Base', 'Peso', 'Ponderado']]
                for criterio, detalle in resultado.detalle_criterios.items():
                    criterios_data.append([
                        criterio,
                        self.formatear_numero(detalle['score_base']),
                        self.formatear_numero(detalle['peso'] * 100, 1) + '%',
                        self.formatear_numero(detalle['score_ponderado'])
                    ])

                tabla_criterios = Table(criterios_data, colWidths=[2 * inch, 1.2 * inch, 1 * inch, 1.2 * inch])
                tabla_criterios.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                ]))

                story.append(tabla_criterios)

                # Observaciones
                if resultado.observaciones:
                    story.append(Spacer(1, 0.1 * inch))
                    story.append(Paragraph("<b>Observaciones:</b>", styles['Normal']))
                    for obs in resultado.observaciones:
                        story.append(Paragraph(f"• {obs}", styles['Normal']))

        # Generar PDF
        doc.build(story)
        return output.getvalue()
