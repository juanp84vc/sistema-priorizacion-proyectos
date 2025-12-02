"""
Módulo para exportar evaluación de un proyecto individual a diferentes formatos.
Diseñado para presentaciones ejecutivas a juntas directivas.

Soporta: Word, PDF, Excel y Resumen Ejecutivo (1 página)

Actualizado Dic 2025: Integración con ExplicadorCriterios para explicaciones
detalladas con fórmulas paso a paso.
"""
import io
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
import pandas as pd

# Agregar src al path para imports locales
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from reportes.explicaciones_criterios import ExplicadorCriterios
except ImportError:
    ExplicadorCriterios = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
except ImportError:
    SimpleDocTemplate = None


class ExportadorProyectoIndividual:
    """
    Exporta la evaluación completa de un proyecto individual.
    
    Diseñado para generar documentos ejecutivos apropiados para
    presentaciones a juntas directivas y seguimiento de proyectos.
    """
    
    def __init__(self, proyecto, resultado, datos_basicos: Dict[str, Any]):
        """
        Inicializa el exportador con los datos del proyecto.
        
        Args:
            proyecto: Objeto ProyectoSocial evaluado
            resultado: ResultadoScoring del motor Arquitectura C
            datos_basicos: Diccionario con datos básicos del formulario
        """
        self.proyecto = proyecto
        self.resultado = resultado
        self.datos_basicos = datos_basicos
        self.fecha_evaluacion = datetime.now().strftime("%d/%m/%Y %H:%M")
    
    @staticmethod
    def formatear_numero(numero: float, decimales: int = 2) -> str:
        """
        Formatea un número con punto para miles y coma para decimales.
        
        Args:
            numero: Número a formatear
            decimales: Cantidad de decimales a mostrar
        
        Returns:
            str: Número formateado (ej: 1.234.567,89)
        """
        if numero is None:
            return "0,00"
        
        formato = f"{{:,.{decimales}f}}"
        numero_formateado = formato.format(numero)
        
        # Intercambiar punto y coma
        numero_formateado = numero_formateado.replace(",", "TEMP")
        numero_formateado = numero_formateado.replace(".", ",")
        numero_formateado = numero_formateado.replace("TEMP", ".")
        
        return numero_formateado
    
    def _generar_analisis_sroi(self) -> str:
        """
        Genera análisis detallado del SROI según Arquitectura C.
        """
        sroi = self.proyecto.indicadores_impacto.get('sroi', 0)
        
        if sroi >= 4.0:
            interpretacion = "EXCELENTE - El proyecto genera un retorno social muy superior a la inversión. Por cada peso invertido, se generan más de 4 pesos de valor social."
            recomendacion = "Proyecto altamente recomendado. El retorno social justifica ampliamente la inversión."
        elif sroi >= 3.0:
            interpretacion = "MUY BUENO - El proyecto genera un retorno social significativo. Por cada peso invertido, se generan entre 3 y 4 pesos de valor social."
            recomendacion = "Proyecto recomendado. El retorno social es sólido y supera ampliamente el umbral mínimo."
        elif sroi >= 2.0:
            interpretacion = "BUENO - El proyecto genera un retorno social positivo. Por cada peso invertido, se generan entre 2 y 3 pesos de valor social."
            recomendacion = "Proyecto viable. El retorno social es aceptable aunque podría mejorarse."
        elif sroi >= 1.0:
            interpretacion = "MARGINAL - El proyecto apenas genera valor social adicional. El retorno es cercano al punto de equilibrio."
            recomendacion = "Proyecto requiere optimización. Considerar ajustes para mejorar el retorno social."
        else:
            interpretacion = "INSUFICIENTE - El proyecto no genera valor social neto positivo. La inversión supera el retorno social."
            recomendacion = "Proyecto NO recomendado. Requiere rediseño fundamental o rechazo."
        
        return f"""
Interpretación del SROI ({sroi:.2f}):
{interpretacion}

Recomendación:
{recomendacion}

Costo-Beneficio Social:
- Inversión: ${self.formatear_numero(self.proyecto.presupuesto_total, 0)} COP
- Valor Social Generado: ${self.formatear_numero(self.proyecto.presupuesto_total * sroi, 0)} COP
- Beneficio Neto: ${self.formatear_numero(self.proyecto.presupuesto_total * (sroi - 1), 0)} COP
"""
    
    def _generar_analisis_stakeholders(self) -> str:
        """
        Genera análisis detallado de stakeholders según Arquitectura C.
        """
        pertinencia = self.proyecto.pertinencia_operacional
        relacionamiento = self.proyecto.mejora_relacionamiento
        corredor = self.proyecto.en_corredor_transmision
        num_stakeholders = len(self.proyecto.stakeholders_involucrados)
        
        analisis = f"""
Análisis de Impacto en Stakeholders:

1. Pertinencia Operacional ({pertinencia}/5):
"""
        if pertinencia >= 4:
            analisis += "   - ALTA: El proyecto está altamente alineado con las necesidades operacionales de GEB.\n"
        elif pertinencia >= 3:
            analisis += "   - MEDIA: El proyecto tiene alineación moderada con las necesidades operacionales.\n"
        else:
            analisis += "   - BAJA: El proyecto tiene baja alineación con las necesidades operacionales.\n"
        
        analisis += f"\n2. Mejora de Relacionamiento ({relacionamiento}/5):\n"
        if relacionamiento >= 4:
            analisis += "   - ALTA: El proyecto mejora significativamente las relaciones con stakeholders clave.\n"
        elif relacionamiento >= 3:
            analisis += "   - MEDIA: El proyecto tiene impacto moderado en el relacionamiento.\n"
        else:
            analisis += "   - BAJA: El proyecto tiene impacto limitado en el relacionamiento.\n"
        
        analisis += f"\n3. Ubicación Estratégica:\n"
        if corredor:
            analisis += "   - ✅ VENTAJA: Proyecto ubicado en corredor de transmisión, maximizando impacto operacional.\n"
        else:
            analisis += "   - ⚠️ LIMITACIÓN: Proyecto fuera de corredor de transmisión, impacto operacional reducido.\n"
        
        analisis += f"\n4. Alcance de Stakeholders ({num_stakeholders} tipos involucrados):\n"
        if num_stakeholders >= 4:
            analisis += "   - AMPLIO: Proyecto involucra múltiples tipos de stakeholders, maximizando impacto social.\n"
        elif num_stakeholders >= 2:
            analisis += "   - MODERADO: Proyecto involucra varios stakeholders clave.\n"
        else:
            analisis += "   - LIMITADO: Proyecto involucra pocos stakeholders, impacto social acotado.\n"
        
        return analisis
    
    def _generar_analisis_pdet(self) -> str:
        """
        Genera análisis detallado de elegibilidad PDET según Arquitectura C.
        """
        if not self.proyecto.tiene_municipios_pdet:
            return """
Análisis de Elegibilidad PDET:

❌ MUNICIPIO NO-PDET

Implicaciones:
- El proyecto NO es elegible para mecanismo de Obras por Impuestos
- Score de Probabilidad de Aprobación: 0/100
- Impacto significativo en score final (-20 puntos potenciales)

Recomendación:
- Considerar reubicación a municipio PDET si es viable
- Explorar mecanismos alternativos de financiación
- Evaluar si el alto impacto social justifica la inversión directa
"""
        
        sector = self.proyecto.sectores[0] if self.proyecto.sectores else 'No especificado'
        puntaje_sectorial = self.proyecto.puntaje_sectorial_max or 0
        
        analisis = f"""
Análisis de Elegibilidad PDET:

✅ MUNICIPIO PDET CONFIRMADO

Ventajas:
- Elegible para mecanismo de Obras por Impuestos
- Acceso a beneficios tributarios para donantes
- Alineación con política pública de construcción de paz

Análisis Sectorial:
- Sector: {sector.title()}
- Puntaje Sectorial PDET: {puntaje_sectorial}/10
"""
        
        if puntaje_sectorial >= 8:
            analisis += "\nPrioridad Sectorial: MÁXIMA\n"
            analisis += "El sector es altamente prioritario según lineamientos PDET. Probabilidad de aprobación muy alta.\n"
        elif puntaje_sectorial >= 6:
            analisis += "\nPrioridad Sectorial: ALTA\n"
            analisis += "El sector es prioritario según lineamientos PDET. Buena probabilidad de aprobación.\n"
        elif puntaje_sectorial >= 4:
            analisis += "\nPrioridad Sectorial: MEDIA\n"
            analisis += "El sector tiene prioridad moderada. Probabilidad de aprobación aceptable.\n"
        else:
            analisis += "\nPrioridad Sectorial: BAJA\n"
            analisis += "El sector tiene baja prioridad. Probabilidad de aprobación reducida.\n"
        
        return analisis
    
    def _generar_analisis_riesgos(self) -> str:
        """
        Genera análisis detallado de riesgos con estrategias de mitigación.
        """
        riesgo_tecnico = self.proyecto.riesgo_tecnico_probabilidad * self.proyecto.riesgo_tecnico_impacto
        riesgo_social = self.proyecto.riesgo_social_probabilidad * self.proyecto.riesgo_social_impacto
        riesgo_financiero = self.proyecto.riesgo_financiero_probabilidad * self.proyecto.riesgo_financiero_impacto
        riesgo_regulatorio = self.proyecto.riesgo_regulatorio_probabilidad * self.proyecto.riesgo_regulatorio_impacto
        
        riesgo_total = (riesgo_tecnico + riesgo_social + riesgo_financiero + riesgo_regulatorio) / 4
        
        analisis = f"""
Análisis de Riesgos y Estrategias de Mitigación:

Perfil de Riesgo General: {riesgo_total:.1f}/25
"""
        
        if riesgo_total <= 6:
            analisis += "Clasificación: BAJO RIESGO - Proyecto con alta viabilidad\n\n"
        elif riesgo_total <= 12:
            analisis += "Clasificación: RIESGO MODERADO - Requiere plan de mitigación estándar\n\n"
        else:
            analisis += "Clasificación: ALTO RIESGO - Requiere plan de mitigación robusto\n\n"
        
        # Análisis por tipo de riesgo
        analisis += f"1. Riesgo Técnico: {riesgo_tecnico}/25\n"
        if riesgo_tecnico >= 12:
            analisis += "   - CRÍTICO: Requiere evaluación técnica exhaustiva\n"
            analisis += "   - Mitigación: Contratar consultoría especializada, pruebas piloto\n"
        elif riesgo_tecnico >= 6:
            analisis += "   - MODERADO: Requiere supervisión técnica adecuada\n"
            analisis += "   - Mitigación: Equipo técnico calificado, seguimiento regular\n"
        else:
            analisis += "   - BAJO: Tecnología probada y equipo capacitado\n"
            analisis += "   - Mitigación: Seguimiento estándar\n"
        
        analisis += f"\n2. Riesgo Social: {riesgo_social}/25\n"
        if riesgo_social >= 12:
            analisis += "   - CRÍTICO: Requiere estrategia robusta de relacionamiento comunitario\n"
            analisis += "   - Mitigación: Consulta previa, acuerdos comunitarios, mediación\n"
        elif riesgo_social >= 6:
            analisis += "   - MODERADO: Requiere gestión activa de stakeholders\n"
            analisis += "   - Mitigación: Comunicación constante, participación comunitaria\n"
        else:
            analisis += "   - BAJO: Buena aceptación comunitaria esperada\n"
            analisis += "   - Mitigación: Comunicación transparente\n"
        
        analisis += f"\n3. Riesgo Financiero: {riesgo_financiero}/25\n"
        if riesgo_financiero >= 12:
            analisis += "   - CRÍTICO: Requiere estructura financiera robusta\n"
            analisis += "   - Mitigación: Diversificación de fuentes, garantías, reservas\n"
        elif riesgo_financiero >= 6:
            analisis += "   - MODERADO: Requiere seguimiento financiero estricto\n"
            analisis += "   - Mitigación: Presupuesto contingente, monitoreo mensual\n"
        else:
            analisis += "   - BAJO: Estructura financiera sólida\n"
            analisis += "   - Mitigación: Seguimiento estándar\n"
        
        analisis += f"\n4. Riesgo Regulatorio: {riesgo_regulatorio}/25\n"
        if riesgo_regulatorio >= 12:
            analisis += "   - CRÍTICO: Requiere asesoría legal especializada\n"
            analisis += "   - Mitigación: Revisión legal exhaustiva, permisos anticipados\n"
        elif riesgo_regulatorio >= 6:
            analisis += "   - MODERADO: Requiere cumplimiento normativo riguroso\n"
            analisis += "   - Mitigación: Asesoría legal, gestión de permisos\n"
        else:
            analisis += "   - BAJO: Marco regulatorio claro\n"
            analisis += "   - Mitigación: Cumplimiento normativo estándar\n"
        
        return analisis

    def _generar_explicacion_detallada_sroi(self) -> Dict[str, Any]:
        """
        Genera explicación detallada del SROI con fórmulas paso a paso.
        Usa el módulo ExplicadorCriterios si está disponible.
        """
        if ExplicadorCriterios is None:
            return {"error": "Módulo ExplicadorCriterios no disponible"}

        sroi = self.proyecto.indicadores_impacto.get('sroi', 0)
        return ExplicadorCriterios.explicar_sroi(
            sroi_valor=sroi,
            score=self.resultado.score_sroi,
            contribucion=self.resultado.contribucion_sroi
        )

    def _generar_explicacion_detallada_stakeholders(self) -> Dict[str, Any]:
        """
        Genera explicación detallada de Stakeholders con fórmulas.
        """
        if ExplicadorCriterios is None:
            return {"error": "Módulo ExplicadorCriterios no disponible"}

        # Calcular scores de componentes (aproximados basados en inputs)
        escala_pert = {5: 100, 4: 85, 3: 65, 2: 40, 1: 20}
        escala_rel = {5: 100, 4: 85, 3: 65, 2: 40, 1: 20}

        score_pertinencia = escala_pert.get(self.proyecto.pertinencia_operacional, 50)
        score_relacionamiento = escala_rel.get(self.proyecto.mejora_relacionamiento, 50)

        # Alcance territorial
        score_alcance = 50
        if self.proyecto.en_corredor_transmision:
            score_alcance += 30
        if len(self.proyecto.municipios or []) > 1:
            score_alcance += 20
        score_alcance = min(score_alcance, 100)

        # Tipo stakeholders
        puntajes_stak = {
            'autoridades_locales': 25, 'lideres_comunitarios': 20,
            'comunidades_indigenas': 25, 'organizaciones_sociales': 15,
            'sector_privado': 10, 'academia': 10, 'medios_comunicacion': 5
        }
        puntaje_total = sum(puntajes_stak.get(s, 0) for s in (self.proyecto.stakeholders_involucrados or []))
        score_tipo = min((puntaje_total / 110) * 100, 100) if puntaje_total > 0 else 50

        return ExplicadorCriterios.explicar_stakeholders(
            pertinencia=self.proyecto.pertinencia_operacional,
            relacionamiento=self.proyecto.mejora_relacionamiento,
            score_pertinencia=score_pertinencia,
            score_relacionamiento=score_relacionamiento,
            score_alcance=score_alcance,
            score_tipo=score_tipo,
            score_total=self.resultado.score_stakeholders,
            contribucion=self.resultado.contribucion_stakeholders,
            en_corredor=self.proyecto.en_corredor_transmision,
            stakeholders_lista=self.proyecto.stakeholders_involucrados or [],
            num_municipios=len(self.proyecto.municipios or [])
        )

    def _generar_explicacion_detallada_probabilidad(self) -> Dict[str, Any]:
        """
        Genera explicación detallada de Probabilidad PDET con fórmulas.
        """
        if ExplicadorCriterios is None:
            return {"error": "Módulo ExplicadorCriterios no disponible"}

        municipio = self.datos_basicos.get('municipio', 'N/A')
        departamento = self.datos_basicos.get('departamento', 'N/A')
        sector = self.proyecto.sectores[0] if self.proyecto.sectores else 'No especificado'

        return ExplicadorCriterios.explicar_probabilidad(
            municipio=municipio,
            departamento=departamento,
            es_pdet=self.proyecto.tiene_municipios_pdet,
            sector=sector,
            puntaje_sectorial=self.proyecto.puntaje_sectorial_max or 0,
            score=self.resultado.score_probabilidad,
            contribucion=self.resultado.contribucion_probabilidad
        )

    def _generar_explicacion_detallada_riesgos(self) -> Dict[str, Any]:
        """
        Genera explicación detallada de Riesgos con fórmulas.
        """
        if ExplicadorCriterios is None:
            return {"error": "Módulo ExplicadorCriterios no disponible"}

        # Calcular scores individuales de riesgos
        def calc_score_riesgo(prob, imp):
            severidad = prob * imp
            return max(0, 100 - (severidad * 4))

        score_tecnico = calc_score_riesgo(
            self.proyecto.riesgo_tecnico_probabilidad,
            self.proyecto.riesgo_tecnico_impacto
        )
        score_social = calc_score_riesgo(
            self.proyecto.riesgo_social_probabilidad,
            self.proyecto.riesgo_social_impacto
        )
        score_financiero = calc_score_riesgo(
            self.proyecto.riesgo_financiero_probabilidad,
            self.proyecto.riesgo_financiero_impacto
        )
        score_regulatorio = calc_score_riesgo(
            self.proyecto.riesgo_regulatorio_probabilidad,
            self.proyecto.riesgo_regulatorio_impacto
        )

        return ExplicadorCriterios.explicar_riesgos(
            riesgo_tecnico=(self.proyecto.riesgo_tecnico_probabilidad, self.proyecto.riesgo_tecnico_impacto),
            riesgo_social=(self.proyecto.riesgo_social_probabilidad, self.proyecto.riesgo_social_impacto),
            riesgo_financiero=(self.proyecto.riesgo_financiero_probabilidad, self.proyecto.riesgo_financiero_impacto),
            riesgo_regulatorio=(self.proyecto.riesgo_regulatorio_probabilidad, self.proyecto.riesgo_regulatorio_impacto),
            score_tecnico=score_tecnico,
            score_social=score_social,
            score_financiero=score_financiero,
            score_regulatorio=score_regulatorio,
            score_total=self.resultado.score_riesgos,
            contribucion=self.resultado.contribucion_riesgos
        )

    def _agregar_seccion_metodologia_word(self, doc) -> None:
        """
        Agrega sección de metodología Arquitectura C al documento Word.
        """
        if ExplicadorCriterios is None:
            return

        doc.add_heading('Metodología de Evaluación: Arquitectura C', 1)

        metodologia = ExplicadorCriterios.generar_resumen_metodologia()

        doc.add_paragraph(metodologia['descripcion'])
        doc.add_paragraph()

        doc.add_heading('Fórmula de Cálculo', 2)
        doc.add_paragraph(metodologia['formula_general'])
        doc.add_paragraph()

        # Tabla de criterios
        tabla = doc.add_table(rows=5, cols=3)
        tabla.style = 'Light Grid Accent 1'

        headers = tabla.rows[0].cells
        headers[0].text = "Criterio"
        headers[1].text = "Peso"
        headers[2].text = "Descripción"

        for i, criterio in enumerate(metodologia['criterios'], 1):
            row = tabla.rows[i].cells
            row[0].text = criterio['nombre']
            row[1].text = criterio['peso']
            row[2].text = criterio['descripcion']

        doc.add_paragraph()

        # Niveles de prioridad
        doc.add_heading('Niveles de Prioridad', 2)

        tabla_niveles = doc.add_table(rows=6, cols=3)
        tabla_niveles.style = 'Light List Accent 1'

        headers = tabla_niveles.rows[0].cells
        headers[0].text = "Rango Score"
        headers[1].text = "Nivel"
        headers[2].text = "Recomendación"

        for i, nivel in enumerate(metodologia['niveles_prioridad'], 1):
            row = tabla_niveles.rows[i].cells
            row[0].text = nivel['rango']
            row[1].text = nivel['nivel']
            row[2].text = nivel['recomendacion']

    def _agregar_explicacion_criterio_word(self, doc, explicacion: Dict[str, Any]) -> None:
        """
        Agrega una explicación detallada de criterio al documento Word.
        """
        if 'error' in explicacion:
            return

        # Título y definición
        doc.add_paragraph(f"Definición: {explicacion.get('definicion', '')}")
        doc.add_paragraph()

        # Valores y cálculo
        if 'valor_input' in explicacion:
            doc.add_paragraph(f"Valor de entrada: {explicacion['valor_input']}")

        doc.add_paragraph(f"Score obtenido: {explicacion.get('score', 0):.1f}/100")

        if 'formula_aplicada' in explicacion:
            doc.add_paragraph(f"Fórmula aplicada: {explicacion['formula_aplicada']}")

        if 'calculo_contribucion' in explicacion:
            doc.add_paragraph(f"Cálculo: {explicacion['calculo_contribucion']}")

        doc.add_paragraph()

        # Interpretación
        if 'interpretacion' in explicacion:
            p = doc.add_paragraph()
            p.add_run("Interpretación: ").bold = True
            p.add_run(explicacion['interpretacion'])

        # Recomendación si existe
        if 'recomendacion' in explicacion:
            p = doc.add_paragraph()
            p.add_run("Recomendación: ").bold = True
            p.add_run(explicacion['recomendacion'])

        doc.add_paragraph()

    def _agregar_explicacion_criterio_pdf(self, story, styles, explicacion: Dict[str, Any], titulo: str) -> None:
        """
        Agrega una explicación detallada de criterio al PDF.
        """
        if 'error' in explicacion:
            return

        heading_style = ParagraphStyle(
            'CriterioHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=8
        )

        story.append(Paragraph(titulo, heading_style))

        # Definición
        if 'definicion' in explicacion:
            story.append(Paragraph(f"<b>Definición:</b> {explicacion['definicion']}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

        # Valores y cálculo
        if 'valor_input' in explicacion:
            story.append(Paragraph(f"<b>Valor de entrada:</b> {explicacion['valor_input']}", styles['Normal']))

        story.append(Paragraph(f"<b>Score obtenido:</b> {explicacion.get('score', 0):.1f}/100", styles['Normal']))

        if 'formula_aplicada' in explicacion:
            story.append(Paragraph(f"<b>Fórmula:</b> {explicacion['formula_aplicada']}", styles['Normal']))

        if 'calculo_contribucion' in explicacion:
            story.append(Paragraph(f"<b>Cálculo:</b> {explicacion['calculo_contribucion']}", styles['Normal']))

        story.append(Spacer(1, 0.1*inch))

        # Interpretación
        if 'interpretacion' in explicacion:
            story.append(Paragraph(f"<b>Interpretación:</b> {explicacion['interpretacion']}", styles['Normal']))

        # Recomendación
        if 'recomendacion' in explicacion:
            story.append(Paragraph(f"<b>Recomendación:</b> {explicacion['recomendacion']}", styles['Normal']))

        story.append(Spacer(1, 0.15*inch))

    def _agregar_seccion_metodologia_pdf(self, story, styles) -> None:
        """
        Agrega sección de metodología al PDF.
        """
        if ExplicadorCriterios is None:
            return

        heading_style = ParagraphStyle(
            'MetodologiaHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=12
        )

        story.append(Paragraph("Metodología de Evaluación: Arquitectura C", heading_style))

        metodologia = ExplicadorCriterios.generar_resumen_metodologia()

        story.append(Paragraph(metodologia['descripcion'], styles['Normal']))
        story.append(Spacer(1, 0.1*inch))

        story.append(Paragraph("<b>Fórmula General:</b>", styles['Normal']))
        story.append(Paragraph(metodologia['formula_general'], styles['Normal']))
        story.append(Spacer(1, 0.15*inch))

        # Tabla de criterios
        criterios_met = [['Criterio', 'Peso', 'Descripción']]
        for criterio in metodologia['criterios']:
            criterios_met.append([criterio['nombre'], criterio['peso'], criterio['descripcion']])

        criterios_table = Table(criterios_met, colWidths=[2*inch, 0.8*inch, 3.5*inch])
        criterios_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f2f6')])
        ]))

        story.append(criterios_table)
        story.append(Spacer(1, 0.2*inch))

        # Niveles de prioridad
        story.append(Paragraph("<b>Niveles de Prioridad:</b>", styles['Normal']))
        story.append(Spacer(1, 0.05*inch))

        niveles_data = [['Rango', 'Nivel', 'Recomendación']]
        for nivel in metodologia['niveles_prioridad']:
            niveles_data.append([nivel['rango'], nivel['nivel'], nivel['recomendacion']])

        niveles_table = Table(niveles_data, colWidths=[1.5*inch, 1.5*inch, 3.3*inch])
        niveles_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f2f6')])
        ]))

        story.append(niveles_table)

    def exportar_word(self) -> bytes:
        """
        Exporta a Word con formato ejecutivo profesional.

        Returns:
            bytes: Contenido del archivo Word
        """
        if Document is None:
            raise ImportError("python-docx no está instalado")
        
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # ========== PORTADA ==========
        # Título principal
        titulo = doc.add_heading(self.proyecto.nombre, 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitulo = doc.add_paragraph(f"Evaluación de Proyecto - Arquitectura C")
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitulo.runs[0].font.size = Pt(14)
        subtitulo.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        
        # Score prominente
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = score_para.add_run(f"{self.resultado.score_total:.1f}/100")
        score_run.font.size = Pt(48)
        score_run.font.bold = True
        
        # Color según nivel
        if self.resultado.score_total >= 85:
            score_run.font.color.rgb = RGBColor(34, 197, 94)  # Verde
        elif self.resultado.score_total >= 70:
            score_run.font.color.rgb = RGBColor(234, 179, 8)  # Amarillo
        else:
            score_run.font.color.rgb = RGBColor(239, 68, 68)  # Rojo
        
        # Nivel de prioridad
        nivel_para = doc.add_paragraph(self.resultado.nivel_prioridad)
        nivel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nivel_para.runs[0].font.size = Pt(18)
        nivel_para.runs[0].font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Información de evaluación
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_table.rows[0].cells[0].text = "Organización:"
        info_table.rows[0].cells[1].text = self.proyecto.organizacion
        info_table.rows[1].cells[0].text = "Ubicación:"
        info_table.rows[1].cells[1].text = f"{self.datos_basicos['municipio']}, {self.datos_basicos['departamento']}"
        info_table.rows[2].cells[0].text = "Fecha de Evaluación:"
        info_table.rows[2].cells[1].text = self.fecha_evaluacion
        info_table.rows[3].cells[0].text = "Sistema:"
        info_table.rows[3].cells[1].text = "Motor Arquitectura C v1.0"
        
        doc.add_page_break()
        
        # ========== RESUMEN EJECUTIVO ==========
        doc.add_heading('Resumen Ejecutivo', 1)
        
        # Datos clave
        doc.add_heading('Datos Clave del Proyecto', 2)
        
        datos_table = doc.add_table(rows=6, cols=2)
        datos_table.style = 'Light List Accent 1'
        
        datos_table.rows[0].cells[0].text = "Presupuesto Total:"
        datos_table.rows[0].cells[1].text = f"${self.formatear_numero(self.proyecto.presupuesto_total, 0)} COP"
        
        datos_table.rows[1].cells[0].text = "Beneficiarios Directos:"
        datos_table.rows[1].cells[1].text = self.formatear_numero(self.proyecto.beneficiarios_directos, 0)
        
        datos_table.rows[2].cells[0].text = "Beneficiarios Indirectos:"
        datos_table.rows[2].cells[1].text = self.formatear_numero(self.proyecto.beneficiarios_indirectos, 0)
        
        datos_table.rows[3].cells[0].text = "Duración Estimada:"
        datos_table.rows[3].cells[1].text = f"{self.proyecto.duracion_estimada_meses} meses"
        
        datos_table.rows[4].cells[0].text = "Municipio PDET:"
        datos_table.rows[4].cells[1].text = "✅ Sí" if self.proyecto.tiene_municipios_pdet else "❌ No"
        
        costo_beneficiario = self.proyecto.presupuesto_total / self.proyecto.beneficiarios_directos if self.proyecto.beneficiarios_directos > 0 else 0
        datos_table.rows[5].cells[0].text = "Costo por Beneficiario:"
        datos_table.rows[5].cells[1].text = f"${self.formatear_numero(costo_beneficiario, 0)} COP"
        
        doc.add_paragraph()
        
        # Evaluación por criterios
        doc.add_heading('Evaluación por Criterios (Arquitectura C)', 2)
        
        criterios_table = doc.add_table(rows=5, cols=4)
        criterios_table.style = 'Medium Shading 1 Accent 1'
        
        # Headers
        headers = criterios_table.rows[0].cells
        headers[0].text = "Criterio"
        headers[1].text = "Peso"
        headers[2].text = "Score"
        headers[3].text = "Contribución"
        
        # SROI
        row1 = criterios_table.rows[1].cells
        row1[0].text = "SROI (Retorno Social)"
        row1[1].text = "40%"
        row1[2].text = f"{self.resultado.score_sroi:.0f}/100"
        row1[3].text = f"{self.resultado.contribucion_sroi:.1f} pts"
        
        # Stakeholders
        row2 = criterios_table.rows[2].cells
        row2[0].text = "Stakeholders"
        row2[1].text = "25%"
        row2[2].text = f"{self.resultado.score_stakeholders:.0f}/100"
        row2[3].text = f"{self.resultado.contribucion_stakeholders:.1f} pts"
        
        # Probabilidad
        row3 = criterios_table.rows[3].cells
        row3[0].text = "Probabilidad Aprobación"
        row3[1].text = "20%"
        row3[2].text = f"{self.resultado.score_probabilidad:.0f}/100"
        row3[3].text = f"{self.resultado.contribucion_probabilidad:.1f} pts"
        
        # Riesgos
        row4 = criterios_table.rows[4].cells
        row4[0].text = "Riesgos"
        row4[1].text = "15%"
        row4[2].text = f"{self.resultado.score_riesgos:.0f}/100"
        row4[3].text = f"{self.resultado.contribucion_riesgos:.1f} pts"
        
        doc.add_paragraph()
        
        # Alertas y recomendaciones
        if self.resultado.alertas:
            doc.add_heading('Alertas', 2)
            for alerta in self.resultado.alertas:
                p = doc.add_paragraph(alerta, style='List Bullet')
                if "RECHAZADO" in alerta or "CRÍTICO" in alerta:
                    p.runs[0].font.color.rgb = RGBColor(239, 68, 68)
                    p.runs[0].font.bold = True
        
        if self.resultado.recomendaciones:
            doc.add_heading('Recomendaciones', 2)
            for rec in self.resultado.recomendaciones:
                doc.add_paragraph(rec, style='List Bullet')
        
        doc.add_page_break()
        
        # ========== ANÁLISIS DETALLADO ==========
        doc.add_heading('Análisis Detallado por Criterio', 1)
        
        # SROI
        doc.add_heading('1. SROI - Social Return on Investment (40%)', 2)
        doc.add_paragraph(f"Valor SROI: {self.proyecto.indicadores_impacto.get('sroi', 0):.2f}")
        doc.add_paragraph(f"Score: {self.resultado.score_sroi:.0f}/100")
        doc.add_paragraph(f"Contribución al score final: {self.resultado.contribucion_sroi:.1f} puntos")
        
        # Stakeholders
        doc.add_heading('2. Stakeholders (25%)', 2)
        doc.add_paragraph(f"Pertinencia Operacional: {self.proyecto.pertinencia_operacional}/5")
        doc.add_paragraph(f"Mejora Relacionamiento: {self.proyecto.mejora_relacionamiento}/5")
        doc.add_paragraph(f"En Corredor de Transmisión: {'Sí' if self.proyecto.en_corredor_transmision else 'No'}")
        doc.add_paragraph(f"Stakeholders Involucrados: {len(self.proyecto.stakeholders_involucrados)}")
        doc.add_paragraph(f"Score: {self.resultado.score_stakeholders:.0f}/100")
        
        # Probabilidad
        doc.add_heading('3. Probabilidad de Aprobación (20%)', 2)
        if self.proyecto.tiene_municipios_pdet:
            doc.add_paragraph(f"Municipio PDET: ✅ Sí")
            if self.proyecto.sectores:
                doc.add_paragraph(f"Sector: {self.proyecto.sectores[0].title()}")
            doc.add_paragraph(f"Puntaje Sectorial: {self.proyecto.puntaje_sectorial_max}/10")
        else:
            doc.add_paragraph(f"Municipio PDET: ❌ No")
            doc.add_paragraph("No elegible para Obras por Impuestos")
        doc.add_paragraph(f"Score: {self.resultado.score_probabilidad:.0f}/100")
        
        # Riesgos
        doc.add_heading('4. Evaluación de Riesgos (15%)', 2)
        doc.add_paragraph(f"Riesgo Técnico: {self.proyecto.riesgo_tecnico_probabilidad}×{self.proyecto.riesgo_tecnico_impacto} = {self.proyecto.riesgo_tecnico_probabilidad * self.proyecto.riesgo_tecnico_impacto}")
        doc.add_paragraph(f"Riesgo Social: {self.proyecto.riesgo_social_probabilidad}×{self.proyecto.riesgo_social_impacto} = {self.proyecto.riesgo_social_probabilidad * self.proyecto.riesgo_social_impacto}")
        doc.add_paragraph(f"Riesgo Financiero: {self.proyecto.riesgo_financiero_probabilidad}×{self.proyecto.riesgo_financiero_impacto} = {self.proyecto.riesgo_financiero_probabilidad * self.proyecto.riesgo_financiero_impacto}")
        doc.add_paragraph(f"Riesgo Regulatorio: {self.proyecto.riesgo_regulatorio_probabilidad}×{self.proyecto.riesgo_regulatorio_impacto} = {self.proyecto.riesgo_regulatorio_probabilidad * self.proyecto.riesgo_regulatorio_impacto}")
        doc.add_paragraph(f"Score: {self.resultado.score_riesgos:.0f}/100")
        
        doc.add_page_break()

        # ========== METODOLOGÍA ARQUITECTURA C ==========
        self._agregar_seccion_metodologia_word(doc)

        doc.add_page_break()

        # ========== ANÁLISIS DETALLADO POR CRITERIO ==========
        doc.add_heading('Análisis Detallado por Criterio', 1)
        doc.add_paragraph(
            "A continuación se presenta el análisis detallado de cada criterio, "
            "incluyendo los valores de entrada, las fórmulas aplicadas, y la interpretación "
            "de los resultados. Esto permite al Comité Evaluador comprender exactamente "
            "cómo se calculó cada componente del score final."
        )
        doc.add_paragraph()

        # ========== 1. SROI (40%) ==========
        doc.add_heading('1. Retorno Social de la Inversión (SROI) - Peso: 40%', 2)
        explicacion_sroi = self._generar_explicacion_detallada_sroi()
        if 'error' not in explicacion_sroi:
            self._agregar_explicacion_criterio_word(doc, explicacion_sroi)

            # Añadir escala de conversión
            if 'escala_conversion' in explicacion_sroi:
                doc.add_heading('Escala de Conversión SROI → Score', 3)
                tabla_sroi = doc.add_table(rows=5, cols=3)
                tabla_sroi.style = 'Light List Accent 1'

                headers = tabla_sroi.rows[0].cells
                headers[0].text = "Rango SROI"
                headers[1].text = "Score"
                headers[2].text = "Nivel"

                for i, escala in enumerate(explicacion_sroi['escala_conversion'], 1):
                    row = tabla_sroi.rows[i].cells
                    row[0].text = escala['rango']
                    row[1].text = str(escala['score'])
                    row[2].text = escala['nivel']

                doc.add_paragraph()
        else:
            # Fallback al análisis anterior
            analisis_sroi = self._generar_analisis_sroi()
            for linea in analisis_sroi.split('\n'):
                if linea.strip():
                    doc.add_paragraph(linea)

        # ========== 2. STAKEHOLDERS (25%) ==========
        doc.add_heading('2. Stakeholders y Pertinencia Operacional - Peso: 25%', 2)
        explicacion_stak = self._generar_explicacion_detallada_stakeholders()
        if 'error' not in explicacion_stak:
            self._agregar_explicacion_criterio_word(doc, explicacion_stak)

            # Desglose de componentes
            if 'componentes' in explicacion_stak:
                doc.add_heading('Desglose de Componentes', 3)
                tabla_stak = doc.add_table(rows=5, cols=4)
                tabla_stak.style = 'Light List Accent 1'

                headers = tabla_stak.rows[0].cells
                headers[0].text = "Componente"
                headers[1].text = "Peso"
                headers[2].text = "Score"
                headers[3].text = "Contribución"

                for i, comp in enumerate(explicacion_stak['componentes'], 1):
                    row = tabla_stak.rows[i].cells
                    row[0].text = comp['nombre']
                    row[1].text = comp['peso']
                    row[2].text = f"{comp['score']:.1f}"
                    row[3].text = f"{comp['contribucion_parcial']:.2f}"

                doc.add_paragraph()
        else:
            analisis_stakeholders = self._generar_analisis_stakeholders()
            for linea in analisis_stakeholders.split('\n'):
                if linea.strip():
                    doc.add_paragraph(linea)

        # ========== 3. PROBABILIDAD PDET (20%) ==========
        doc.add_heading('3. Probabilidad de Aprobación (PDET) - Peso: 20%', 2)
        explicacion_prob = self._generar_explicacion_detallada_probabilidad()
        if 'error' not in explicacion_prob:
            self._agregar_explicacion_criterio_word(doc, explicacion_prob)

            # Info adicional PDET
            if 'info_pdet' in explicacion_prob:
                doc.add_heading('Información PDET', 3)
                info = explicacion_prob['info_pdet']
                doc.add_paragraph(f"Municipio: {info.get('municipio', 'N/A')}")
                doc.add_paragraph(f"Departamento: {info.get('departamento', 'N/A')}")
                doc.add_paragraph(f"Estado PDET: {'Sí' if info.get('es_pdet') else 'No'}")
                if info.get('sector'):
                    doc.add_paragraph(f"Sector: {info['sector']}")
                doc.add_paragraph()
        else:
            analisis_pdet = self._generar_analisis_pdet()
            for linea in analisis_pdet.split('\n'):
                if linea.strip():
                    doc.add_paragraph(linea)

        # ========== 4. RIESGOS (15%) ==========
        doc.add_heading('4. Evaluación de Riesgos - Peso: 15%', 2)
        explicacion_riesg = self._generar_explicacion_detallada_riesgos()
        if 'error' not in explicacion_riesg:
            self._agregar_explicacion_criterio_word(doc, explicacion_riesg)

            # Matriz de riesgos
            if 'matriz_riesgos' in explicacion_riesg:
                doc.add_heading('Matriz de Riesgos Detallada', 3)
                tabla_riesg = doc.add_table(rows=5, cols=5)
                tabla_riesg.style = 'Light List Accent 1'

                headers = tabla_riesg.rows[0].cells
                headers[0].text = "Tipo de Riesgo"
                headers[1].text = "Prob."
                headers[2].text = "Impacto"
                headers[3].text = "Severidad"
                headers[4].text = "Score"

                for i, riesgo in enumerate(explicacion_riesg['matriz_riesgos'], 1):
                    row = tabla_riesg.rows[i].cells
                    row[0].text = riesgo['tipo']
                    row[1].text = str(riesgo['probabilidad'])
                    row[2].text = str(riesgo['impacto'])
                    row[3].text = str(riesgo['severidad'])
                    row[4].text = f"{riesgo['score']:.1f}"

                doc.add_paragraph()
                doc.add_paragraph(
                    "Nota: Score de riesgo es INVERSO (menor severidad = mayor score). "
                    "Fórmula: Score = 100 - (Severidad × 4)"
                )
        else:
            analisis_riesgos = self._generar_analisis_riesgos()
            for linea in analisis_riesgos.split('\n'):
                if linea.strip():
                    doc.add_paragraph(linea)

        # ========== RESUMEN DE CÁLCULO FINAL ==========
        doc.add_page_break()
        doc.add_heading('Resumen de Cálculo del Score Final', 1)

        doc.add_paragraph(
            "El score final se calcula como la suma ponderada de los scores de cada criterio:"
        )
        doc.add_paragraph()

        # Tabla resumen final
        tabla_final = doc.add_table(rows=6, cols=4)
        tabla_final.style = 'Medium Grid 1 Accent 1'

        headers = tabla_final.rows[0].cells
        headers[0].text = "Criterio"
        headers[1].text = "Peso"
        headers[2].text = "Score"
        headers[3].text = "Contribución"

        criterios_data = [
            ("SROI", "40%", self.resultado.score_sroi, self.resultado.contribucion_sroi),
            ("Stakeholders", "25%", self.resultado.score_stakeholders, self.resultado.contribucion_stakeholders),
            ("Probabilidad", "20%", self.resultado.score_probabilidad, self.resultado.contribucion_probabilidad),
            ("Riesgos", "15%", self.resultado.score_riesgos, self.resultado.contribucion_riesgos),
            ("TOTAL", "100%", "-", self.resultado.score_total)
        ]

        for i, (nombre, peso, score, contrib) in enumerate(criterios_data, 1):
            row = tabla_final.rows[i].cells
            row[0].text = nombre
            row[1].text = peso
            row[2].text = f"{score:.1f}" if isinstance(score, (int, float)) else str(score)
            row[3].text = f"{contrib:.2f}"

        doc.add_paragraph()
        doc.add_paragraph(
            f"Score Final = {self.resultado.contribucion_sroi:.2f} + "
            f"{self.resultado.contribucion_stakeholders:.2f} + "
            f"{self.resultado.contribucion_probabilidad:.2f} + "
            f"{self.resultado.contribucion_riesgos:.2f} = "
            f"{self.resultado.score_total:.2f}"
        )
        
        # Conclusiones y Recomendaciones Finales
        doc.add_page_break()
        doc.add_heading('Conclusiones y Recomendaciones Ejecutivas', 1)
        
        if self.resultado.score_total >= 85:
            doc.add_paragraph("✅ RECOMENDACIÓN: APROBACIÓN INMEDIATA")
            doc.add_paragraph("El proyecto presenta un perfil excepcional en todos los criterios de Arquitectura C. Se recomienda priorización máxima y aprobación expedita.")
        elif self.resultado.score_total >= 70:
            doc.add_paragraph("✅ RECOMENDACIÓN: APROBACIÓN")
            doc.add_paragraph("El proyecto presenta un perfil sólido según Arquitectura C. Se recomienda aprobación con seguimiento estándar.")
        elif self.resultado.score_total >= 50:
            doc.add_paragraph("⚠️ RECOMENDACIÓN: EVALUACIÓN DETALLADA")
            doc.add_paragraph("El proyecto presenta un perfil mixto. Se recomienda evaluación detallada de aspectos críticos antes de decisión final.")
        else:
            doc.add_paragraph("❌ RECOMENDACIÓN: NO APROBAR")
            doc.add_paragraph("El proyecto no cumple con los estándares mínimos de Arquitectura C. Se recomienda rechazo o rediseño fundamental.")
        
        doc.add_paragraph()
        doc.add_paragraph(f"Fecha de Evaluación: {self.fecha_evaluacion}")
        doc.add_paragraph("Sistema: Motor de Priorización Arquitectura C v1.0")
        doc.add_paragraph("Metodología: Scoring ponderado con 4 criterios (SROI 40%, Stakeholders 25%, Probabilidad 20%, Riesgos 15%)")
        
        # Guardar en bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    def exportar_excel(self) -> bytes:
        """
        Exporta a Excel con múltiples hojas: Resumen, Detalles y Criterios.
        
        Returns:
            bytes: Contenido del archivo Excel
        """
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # ========== HOJA 1: RESUMEN ==========
            resumen_data = {
                'Concepto': [
                    'Nombre del Proyecto',
                    'Organización',
                    'Ubicación',
                    'Municipio PDET',
                    'Fecha de Evaluación',
                    '',
                    'SCORE FINAL',
                    'Nivel de Prioridad',
                    '',
                    'Presupuesto Total',
                    'Beneficiarios Directos',
                    'Beneficiarios Indirectos',
                    'Duración (meses)',
                    'Costo por Beneficiario'
                ],
                'Valor': [
                    self.proyecto.nombre,
                    self.proyecto.organizacion,
                    f"{self.datos_basicos['municipio']}, {self.datos_basicos['departamento']}",
                    'Sí' if self.proyecto.tiene_municipios_pdet else 'No',
                    self.fecha_evaluacion,
                    '',
                    f"{self.resultado.score_total:.1f}/100",
                    self.resultado.nivel_prioridad,
                    '',
                    f"${self.formatear_numero(self.proyecto.presupuesto_total, 0)}",
                    self.formatear_numero(self.proyecto.beneficiarios_directos, 0),
                    self.formatear_numero(self.proyecto.beneficiarios_indirectos, 0),
                    self.proyecto.duracion_estimada_meses,
                    f"${self.formatear_numero(self.proyecto.presupuesto_total / self.proyecto.beneficiarios_directos if self.proyecto.beneficiarios_directos > 0 else 0, 0)}"
                ]
            }
            
            df_resumen = pd.DataFrame(resumen_data)
            df_resumen.to_excel(writer, sheet_name='Resumen', index=False)
            
            # ========== HOJA 2: CRITERIOS ==========
            criterios_data = {
                'Criterio': [
                    'SROI (Retorno Social)',
                    'Stakeholders',
                    'Probabilidad de Aprobación',
                    'Riesgos'
                ],
                'Peso (%)': [40, 25, 20, 15],
                'Score (0-100)': [
                    f"{self.resultado.score_sroi:.1f}",
                    f"{self.resultado.score_stakeholders:.1f}",
                    f"{self.resultado.score_probabilidad:.1f}",
                    f"{self.resultado.score_riesgos:.1f}"
                ],
                'Contribución (pts)': [
                    f"{self.resultado.contribucion_sroi:.1f}",
                    f"{self.resultado.contribucion_stakeholders:.1f}",
                    f"{self.resultado.contribucion_probabilidad:.1f}",
                    f"{self.resultado.contribucion_riesgos:.1f}"
                ]
            }
            
            df_criterios = pd.DataFrame(criterios_data)
            df_criterios.to_excel(writer, sheet_name='Criterios', index=False)
            
            # ========== HOJA 3: DETALLES ==========
            detalles_data = {
                'Aspecto': [
                    'SROI - Valor',
                    'SROI - Score',
                    '',
                    'Stakeholders - Pertinencia',
                    'Stakeholders - Relacionamiento',
                    'Stakeholders - Corredor',
                    'Stakeholders - Cantidad',
                    '',
                    'Probabilidad - Municipio PDET',
                    'Probabilidad - Sector',
                    'Probabilidad - Puntaje Sectorial',
                    '',
                    'Riesgos - Técnico',
                    'Riesgos - Social',
                    'Riesgos - Financiero',
                    'Riesgos - Regulatorio'
                ],
                'Valor': [
                    f"{self.proyecto.indicadores_impacto.get('sroi', 0):.2f}",
                    f"{self.resultado.score_sroi:.0f}/100",
                    '',
                    f"{self.proyecto.pertinencia_operacional}/5",
                    f"{self.proyecto.mejora_relacionamiento}/5",
                    'Sí' if self.proyecto.en_corredor_transmision else 'No',
                    len(self.proyecto.stakeholders_involucrados),
                    '',
                    'Sí' if self.proyecto.tiene_municipios_pdet else 'No',
                    self.proyecto.sectores[0].title() if self.proyecto.sectores else 'N/A',
                    f"{self.proyecto.puntaje_sectorial_max}/10" if self.proyecto.puntaje_sectorial_max else 'N/A',
                    '',
                    f"{self.proyecto.riesgo_tecnico_probabilidad}×{self.proyecto.riesgo_tecnico_impacto}={self.proyecto.riesgo_tecnico_probabilidad * self.proyecto.riesgo_tecnico_impacto}",
                    f"{self.proyecto.riesgo_social_probabilidad}×{self.proyecto.riesgo_social_impacto}={self.proyecto.riesgo_social_probabilidad * self.proyecto.riesgo_social_impacto}",
                    f"{self.proyecto.riesgo_financiero_probabilidad}×{self.proyecto.riesgo_financiero_impacto}={self.proyecto.riesgo_financiero_probabilidad * self.proyecto.riesgo_financiero_impacto}",
                    f"{self.proyecto.riesgo_regulatorio_probabilidad}×{self.proyecto.riesgo_regulatorio_impacto}={self.proyecto.riesgo_regulatorio_probabilidad * self.proyecto.riesgo_regulatorio_impacto}"
                ]
            }
            
            df_detalles = pd.DataFrame(detalles_data)
            df_detalles.to_excel(writer, sheet_name='Detalles', index=False)
            
            # ========== HOJA 4: ALERTAS Y RECOMENDACIONES ==========
            alertas_recs = []
            
            if self.resultado.alertas:
                alertas_recs.append({'Tipo': 'ALERTA', 'Mensaje': ''})
                for alerta in self.resultado.alertas:
                    alertas_recs.append({'Tipo': '⚠️', 'Mensaje': alerta})
            
            if self.resultado.recomendaciones:
                alertas_recs.append({'Tipo': '', 'Mensaje': ''})
                alertas_recs.append({'Tipo': 'RECOMENDACIÓN', 'Mensaje': ''})
                for rec in self.resultado.recomendaciones:
                    alertas_recs.append({'Tipo': '💡', 'Mensaje': rec})
            
            if alertas_recs:
                df_alertas = pd.DataFrame(alertas_recs)
                df_alertas.to_excel(writer, sheet_name='Alertas y Recomendaciones', index=False)

            # ========== HOJA 5: METODOLOGÍA ==========
            if ExplicadorCriterios is not None:
                metodologia = ExplicadorCriterios.generar_resumen_metodologia()

                metodologia_data = {
                    'Aspecto': [
                        'Sistema',
                        'Descripción',
                        '',
                        'Fórmula General',
                        '',
                        'CRITERIOS',
                    ],
                    'Detalle': [
                        'Arquitectura C - Motor de Priorización',
                        metodologia['descripcion'],
                        '',
                        metodologia['formula_general'],
                        '',
                        '',
                    ]
                }

                # Agregar criterios
                for criterio in metodologia['criterios']:
                    metodologia_data['Aspecto'].append(f"  {criterio['nombre']}")
                    metodologia_data['Detalle'].append(f"{criterio['peso']} - {criterio['descripcion']}")

                metodologia_data['Aspecto'].append('')
                metodologia_data['Detalle'].append('')
                metodologia_data['Aspecto'].append('NIVELES DE PRIORIDAD')
                metodologia_data['Detalle'].append('')

                for nivel in metodologia['niveles_prioridad']:
                    metodologia_data['Aspecto'].append(f"  {nivel['rango']}")
                    metodologia_data['Detalle'].append(f"{nivel['nivel']} - {nivel['recomendacion']}")

                df_metodologia = pd.DataFrame(metodologia_data)
                df_metodologia.to_excel(writer, sheet_name='Metodología', index=False)

            # ========== HOJA 6: EXPLICACIÓN SROI ==========
            explicacion_sroi = self._generar_explicacion_detallada_sroi()
            if 'error' not in explicacion_sroi:
                sroi_data = {
                    'Campo': [
                        'Criterio',
                        'Peso',
                        'Definición',
                        '',
                        'Valor SROI Ingresado',
                        'Score Obtenido',
                        'Contribución al Total',
                        '',
                        'Fórmula Aplicada',
                        'Interpretación',
                        '',
                        'ESCALA DE CONVERSIÓN',
                        '< 1.0',
                        '1.0 - 1.99',
                        '2.0 - 2.99',
                        '≥ 3.0'
                    ],
                    'Valor': [
                        'SROI - Retorno Social de la Inversión',
                        '40%',
                        explicacion_sroi.get('definicion', ''),
                        '',
                        str(explicacion_sroi.get('valor_input', '')),
                        f"{explicacion_sroi.get('score', 0):.1f}/100",
                        explicacion_sroi.get('calculo_contribucion', ''),
                        '',
                        explicacion_sroi.get('formula_aplicada', ''),
                        explicacion_sroi.get('interpretacion', ''),
                        '',
                        '',
                        'Score: 0 (RECHAZAR)',
                        'Score: 60 (Prioridad Baja)',
                        'Score: 80 (Prioridad Media)',
                        'Score: 95 (Prioridad Alta)'
                    ]
                }
                df_sroi = pd.DataFrame(sroi_data)
                df_sroi.to_excel(writer, sheet_name='Explicación SROI', index=False)

            # ========== HOJA 7: EXPLICACIÓN STAKEHOLDERS ==========
            explicacion_stak = self._generar_explicacion_detallada_stakeholders()
            if 'error' not in explicacion_stak:
                stak_data = {
                    'Campo': [
                        'Criterio',
                        'Peso Total',
                        'Definición',
                        '',
                        'Score Obtenido',
                        'Contribución al Total',
                        '',
                        'COMPONENTES',
                    ],
                    'Valor': [
                        'Stakeholders (Relacionamiento y Pertinencia)',
                        '25%',
                        explicacion_stak.get('definicion', ''),
                        '',
                        f"{explicacion_stak.get('score', 0):.1f}/100",
                        explicacion_stak.get('calculo_contribucion', ''),
                        '',
                        '',
                    ]
                }

                if 'componentes' in explicacion_stak:
                    for comp in explicacion_stak['componentes']:
                        stak_data['Campo'].append(f"  {comp['nombre']} ({comp['peso']})")
                        stak_data['Valor'].append(f"Score: {comp['score']:.1f} → Contrib: {comp['contribucion_parcial']:.2f}")

                stak_data['Campo'].append('')
                stak_data['Valor'].append('')
                stak_data['Campo'].append('Interpretación')
                stak_data['Valor'].append(explicacion_stak.get('interpretacion', ''))

                df_stak = pd.DataFrame(stak_data)
                df_stak.to_excel(writer, sheet_name='Explicación Stakeholders', index=False)

            # ========== HOJA 8: EXPLICACIÓN PROBABILIDAD ==========
            explicacion_prob = self._generar_explicacion_detallada_probabilidad()
            if 'error' not in explicacion_prob:
                prob_data = {
                    'Campo': [
                        'Criterio',
                        'Peso',
                        'Definición',
                        '',
                        'Score Obtenido',
                        'Contribución al Total',
                        '',
                        'INFORMACIÓN PDET',
                        'Municipio',
                        'Departamento',
                        'Es PDET',
                        'Sector',
                        '',
                        'Interpretación'
                    ],
                    'Valor': [
                        'Probabilidad de Aprobación (PDET)',
                        '20%',
                        explicacion_prob.get('definicion', ''),
                        '',
                        f"{explicacion_prob.get('score', 0):.1f}/100",
                        explicacion_prob.get('calculo_contribucion', ''),
                        '',
                        '',
                        explicacion_prob.get('info_pdet', {}).get('municipio', 'N/A') if 'info_pdet' in explicacion_prob else 'N/A',
                        explicacion_prob.get('info_pdet', {}).get('departamento', 'N/A') if 'info_pdet' in explicacion_prob else 'N/A',
                        'Sí' if explicacion_prob.get('info_pdet', {}).get('es_pdet', False) else 'No',
                        explicacion_prob.get('info_pdet', {}).get('sector', 'N/A') if 'info_pdet' in explicacion_prob else 'N/A',
                        '',
                        explicacion_prob.get('interpretacion', '')
                    ]
                }
                df_prob = pd.DataFrame(prob_data)
                df_prob.to_excel(writer, sheet_name='Explicación Probabilidad', index=False)

            # ========== HOJA 9: EXPLICACIÓN RIESGOS ==========
            explicacion_riesg = self._generar_explicacion_detallada_riesgos()
            if 'error' not in explicacion_riesg:
                riesg_data = {
                    'Campo': [
                        'Criterio',
                        'Peso',
                        'Definición',
                        '',
                        'Score Obtenido',
                        'Contribución al Total',
                        '',
                        'Fórmula',
                        '',
                        'MATRIZ DE RIESGOS',
                    ],
                    'Valor': [
                        'Evaluación de Riesgos',
                        '15%',
                        explicacion_riesg.get('definicion', ''),
                        '',
                        f"{explicacion_riesg.get('score', 0):.1f}/100",
                        explicacion_riesg.get('calculo_contribucion', ''),
                        '',
                        'Score = 100 - (Probabilidad × Impacto × 4)',
                        '',
                        'Tipo | Prob | Impacto | Severidad | Score',
                    ]
                }

                if 'matriz_riesgos' in explicacion_riesg:
                    for r in explicacion_riesg['matriz_riesgos']:
                        riesg_data['Campo'].append(f"  {r['tipo']}")
                        riesg_data['Valor'].append(f"{r['probabilidad']} × {r['impacto']} = {r['severidad']} → Score: {r['score']:.1f}")

                riesg_data['Campo'].append('')
                riesg_data['Valor'].append('')
                riesg_data['Campo'].append('Interpretación')
                riesg_data['Valor'].append(explicacion_riesg.get('interpretacion', ''))
                riesg_data['Campo'].append('')
                riesg_data['Valor'].append('')
                riesg_data['Campo'].append('Nota')
                riesg_data['Valor'].append('Score INVERSO: Menor severidad = Mayor score')

                df_riesg = pd.DataFrame(riesg_data)
                df_riesg.to_excel(writer, sheet_name='Explicación Riesgos', index=False)

            # ========== HOJA 10: CÁLCULO FINAL ==========
            calculo_final = {
                'Criterio': ['SROI', 'Stakeholders', 'Probabilidad', 'Riesgos', '', 'TOTAL'],
                'Peso': ['40%', '25%', '20%', '15%', '', '100%'],
                'Score (0-100)': [
                    f"{self.resultado.score_sroi:.1f}",
                    f"{self.resultado.score_stakeholders:.1f}",
                    f"{self.resultado.score_probabilidad:.1f}",
                    f"{self.resultado.score_riesgos:.1f}",
                    '',
                    '-'
                ],
                'Cálculo': [
                    f"{self.resultado.score_sroi:.1f} × 0.40",
                    f"{self.resultado.score_stakeholders:.1f} × 0.25",
                    f"{self.resultado.score_probabilidad:.1f} × 0.20",
                    f"{self.resultado.score_riesgos:.1f} × 0.15",
                    '',
                    f"{self.resultado.contribucion_sroi:.2f} + {self.resultado.contribucion_stakeholders:.2f} + {self.resultado.contribucion_probabilidad:.2f} + {self.resultado.contribucion_riesgos:.2f}"
                ],
                'Contribución': [
                    f"{self.resultado.contribucion_sroi:.2f}",
                    f"{self.resultado.contribucion_stakeholders:.2f}",
                    f"{self.resultado.contribucion_probabilidad:.2f}",
                    f"{self.resultado.contribucion_riesgos:.2f}",
                    '',
                    f"{self.resultado.score_total:.2f}"
                ]
            }
            df_calculo = pd.DataFrame(calculo_final)
            df_calculo.to_excel(writer, sheet_name='Cálculo Final', index=False)

        output.seek(0)
        return output.getvalue()
    
    def exportar_pdf(self) -> bytes:
        """
        Exporta a PDF con gráficos y formato profesional.
        
        Returns:
            bytes: Contenido del archivo PDF
        """
        if SimpleDocTemplate is None:
            raise ImportError("reportlab no está instalado")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12
        )
        
        # ========== PORTADA ==========
        story.append(Spacer(1, 2*inch))
        
        # Título
        story.append(Paragraph(self.proyecto.nombre, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Subtítulo
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=14,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Evaluación de Proyecto - Arquitectura C", subtitle_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Score
        score_style = ParagraphStyle(
            'Score',
            parent=styles['Normal'],
            fontSize=48,
            textColor=colors.HexColor('#22c55e') if self.resultado.score_total >= 85 else colors.HexColor('#eab308') if self.resultado.score_total >= 70 else colors.HexColor('#ef4444'),
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(f"{self.resultado.score_total:.1f}/100", score_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Nivel
        nivel_style = ParagraphStyle(
            'Nivel',
            parent=styles['Normal'],
            fontSize=18,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(self.resultado.nivel_prioridad, nivel_style))
        story.append(Spacer(1, 1*inch))
        
        # Información básica
        info_data = [
            ['Organización:', self.proyecto.organizacion],
            ['Ubicación:', f"{self.datos_basicos['municipio']}, {self.datos_basicos['departamento']}"],
            ['Fecha de Evaluación:', self.fecha_evaluacion],
            ['Sistema:', 'Motor Arquitectura C v1.0']
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f2f6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        
        story.append(info_table)
        story.append(PageBreak())
        
        # ========== RESUMEN EJECUTIVO ==========
        story.append(Paragraph("Resumen Ejecutivo", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Datos clave
        story.append(Paragraph("Datos Clave del Proyecto", styles['Heading3']))
        story.append(Spacer(1, 0.1*inch))
        
        datos_data = [
            ['Concepto', 'Valor'],
            ['Presupuesto Total', f"${self.formatear_numero(self.proyecto.presupuesto_total, 0)} COP"],
            ['Beneficiarios Directos', self.formatear_numero(self.proyecto.beneficiarios_directos, 0)],
            ['Beneficiarios Indirectos', self.formatear_numero(self.proyecto.beneficiarios_indirectos, 0)],
            ['Duración Estimada', f"{self.proyecto.duracion_estimada_meses} meses"],
            ['Municipio PDET', '✓ Sí' if self.proyecto.tiene_municipios_pdet else '✗ No'],
            ['Costo por Beneficiario', f"${self.formatear_numero(self.proyecto.presupuesto_total / self.proyecto.beneficiarios_directos if self.proyecto.beneficiarios_directos > 0 else 0, 0)} COP"]
        ]
        
        datos_table = Table(datos_data, colWidths=[3*inch, 3*inch])
        datos_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f2f6')])
        ]))
        
        story.append(datos_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Evaluación por criterios
        story.append(Paragraph("Evaluación por Criterios (Arquitectura C)", styles['Heading3']))
        story.append(Spacer(1, 0.1*inch))
        
        criterios_data = [
            ['Criterio', 'Peso', 'Score', 'Contribución'],
            ['SROI (Retorno Social)', '40%', f"{self.resultado.score_sroi:.0f}/100", f"{self.resultado.contribucion_sroi:.1f} pts"],
            ['Stakeholders', '25%', f"{self.resultado.score_stakeholders:.0f}/100", f"{self.resultado.contribucion_stakeholders:.1f} pts"],
            ['Probabilidad Aprobación', '20%', f"{self.resultado.score_probabilidad:.0f}/100", f"{self.resultado.contribucion_probabilidad:.1f} pts"],
            ['Riesgos', '15%', f"{self.resultado.score_riesgos:.0f}/100", f"{self.resultado.contribucion_riesgos:.1f} pts"]
        ]
        
        criterios_table = Table(criterios_data, colWidths=[2.5*inch, 1*inch, 1.5*inch, 1.5*inch])
        criterios_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f2f6')])
        ]))
        
        story.append(criterios_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Alertas
        if self.resultado.alertas:
            story.append(Paragraph("Alertas", styles['Heading3']))
            for alerta in self.resultado.alertas:
                story.append(Paragraph(f"• {alerta}", styles['Normal']))
                story.append(Spacer(1, 0.05*inch))
            story.append(Spacer(1, 0.2*inch))
        
        # Recomendaciones
        if self.resultado.recomendaciones:
            story.append(Paragraph("Recomendaciones", styles['Heading3']))
            for rec in self.resultado.recomendaciones:
                story.append(Paragraph(f"• {rec}", styles['Normal']))
                story.append(Spacer(1, 0.05*inch))

        # ========== METODOLOGÍA Y EXPLICACIONES DETALLADAS ==========
        story.append(PageBreak())

        # Sección de Metodología
        self._agregar_seccion_metodologia_pdf(story, styles)

        story.append(PageBreak())

        # Análisis Detallado por Criterio
        detail_heading = ParagraphStyle(
            'DetailHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=12
        )

        story.append(Paragraph("Análisis Detallado por Criterio", detail_heading))
        story.append(Paragraph(
            "A continuación se presenta el desglose de cada criterio con las fórmulas "
            "aplicadas y la interpretación de los resultados.",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2*inch))

        # 1. SROI
        explicacion_sroi = self._generar_explicacion_detallada_sroi()
        self._agregar_explicacion_criterio_pdf(story, styles, explicacion_sroi, "1. SROI - Retorno Social (Peso: 40%)")

        # Tabla de escala SROI
        if 'escala_conversion' in explicacion_sroi:
            sroi_escala = [['Rango SROI', 'Score', 'Nivel']]
            for e in explicacion_sroi['escala_conversion']:
                sroi_escala.append([e['rango'], str(e['score']), e['nivel']])

            sroi_table = Table(sroi_escala, colWidths=[2*inch, 1.5*inch, 2*inch])
            sroi_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            story.append(sroi_table)
            story.append(Spacer(1, 0.2*inch))

        # 2. Stakeholders
        explicacion_stak = self._generar_explicacion_detallada_stakeholders()
        self._agregar_explicacion_criterio_pdf(story, styles, explicacion_stak, "2. Stakeholders (Peso: 25%)")

        # Tabla de componentes Stakeholders
        if 'componentes' in explicacion_stak:
            stak_comp = [['Componente', 'Peso', 'Score', 'Contribución']]
            for c in explicacion_stak['componentes']:
                stak_comp.append([c['nombre'], c['peso'], f"{c['score']:.1f}", f"{c['contribucion_parcial']:.2f}"])

            stak_table = Table(stak_comp, colWidths=[2.5*inch, 1*inch, 1.2*inch, 1.5*inch])
            stak_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            story.append(stak_table)
            story.append(Spacer(1, 0.2*inch))

        # 3. Probabilidad PDET
        explicacion_prob = self._generar_explicacion_detallada_probabilidad()
        self._agregar_explicacion_criterio_pdf(story, styles, explicacion_prob, "3. Probabilidad PDET (Peso: 20%)")

        story.append(PageBreak())

        # 4. Riesgos
        explicacion_riesg = self._generar_explicacion_detallada_riesgos()
        self._agregar_explicacion_criterio_pdf(story, styles, explicacion_riesg, "4. Riesgos (Peso: 15%)")

        # Matriz de Riesgos
        if 'matriz_riesgos' in explicacion_riesg:
            riesg_data = [['Tipo', 'Prob.', 'Impacto', 'Severidad', 'Score']]
            for r in explicacion_riesg['matriz_riesgos']:
                riesg_data.append([r['tipo'], str(r['probabilidad']), str(r['impacto']), str(r['severidad']), f"{r['score']:.1f}"])

            riesg_table = Table(riesg_data, colWidths=[2*inch, 1*inch, 1*inch, 1.2*inch, 1*inch])
            riesg_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            story.append(riesg_table)
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(
                "<i>Nota: Score de riesgo es INVERSO. Fórmula: Score = 100 - (Severidad × 4)</i>",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.2*inch))

        # ========== RESUMEN FINAL DE CÁLCULO ==========
        story.append(Paragraph("Resumen de Cálculo del Score Final", detail_heading))
        story.append(Spacer(1, 0.1*inch))

        # Tabla resumen
        resumen_final = [
            ['Criterio', 'Peso', 'Score', 'Contribución'],
            ['SROI', '40%', f"{self.resultado.score_sroi:.1f}", f"{self.resultado.contribucion_sroi:.2f}"],
            ['Stakeholders', '25%', f"{self.resultado.score_stakeholders:.1f}", f"{self.resultado.contribucion_stakeholders:.2f}"],
            ['Probabilidad', '20%', f"{self.resultado.score_probabilidad:.1f}", f"{self.resultado.contribucion_probabilidad:.2f}"],
            ['Riesgos', '15%', f"{self.resultado.score_riesgos:.1f}", f"{self.resultado.contribucion_riesgos:.2f}"],
            ['TOTAL', '100%', '-', f"{self.resultado.score_total:.2f}"]
        ]

        resumen_table = Table(resumen_final, colWidths=[2*inch, 1.2*inch, 1.5*inch, 1.5*inch])
        resumen_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#22c55e') if self.resultado.score_total >= 70 else colors.HexColor('#eab308')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('TEXTCOLOR', (0, -1), (-1, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))

        story.append(resumen_table)
        story.append(Spacer(1, 0.15*inch))

        story.append(Paragraph(
            f"<b>Score Final</b> = {self.resultado.contribucion_sroi:.2f} + "
            f"{self.resultado.contribucion_stakeholders:.2f} + "
            f"{self.resultado.contribucion_probabilidad:.2f} + "
            f"{self.resultado.contribucion_riesgos:.2f} = "
            f"<b>{self.resultado.score_total:.2f}</b>",
            styles['Normal']
        ))

        # Construir PDF
        doc.build(story)
        buffer.seek(0)

        return buffer.getvalue()
    
    def exportar_resumen_ejecutivo(self) -> bytes:
        """
        Exporta un resumen ejecutivo de 1 página en PDF.
        Ideal para presentaciones rápidas a juntas directivas.
        
        Returns:
            bytes: Contenido del archivo PDF (1 página)
        """
        if SimpleDocTemplate is None:
            raise ImportError("reportlab no está instalado")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Título compacto
        title_style = ParagraphStyle(
            'CompactTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=10,
            alignment=TA_CENTER
        )
        
        story.append(Paragraph(f"{self.proyecto.nombre} - Resumen Ejecutivo", title_style))
        story.append(Spacer(1, 0.1*inch))
        
        # Score prominente
        score_data = [[
            f"SCORE: {self.resultado.score_total:.1f}/100",
            f"NIVEL: {self.resultado.nivel_prioridad}"
        ]]
        
        score_table = Table(score_data, colWidths=[3*inch, 3*inch])
        score_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#22c55e') if self.resultado.score_total >= 85 else colors.HexColor('#eab308') if self.resultado.score_total >= 70 else colors.HexColor('#ef4444')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 14),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12)
        ]))
        
        story.append(score_table)
        story.append(Spacer(1, 0.15*inch))
        
        # Datos clave compactos
        datos_compactos = [
            ['Presupuesto', f"${self.formatear_numero(self.proyecto.presupuesto_total, 0)}", 'Beneficiarios', self.formatear_numero(self.proyecto.beneficiarios_directos, 0)],
            ['Duración', f"{self.proyecto.duracion_estimada_meses} meses", 'PDET', '✓' if self.proyecto.tiene_municipios_pdet else '✗']
        ]
        
        datos_table = Table(datos_compactos, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        datos_table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        
        story.append(datos_table)
        story.append(Spacer(1, 0.15*inch))
        
        # Criterios compactos
        criterios_compactos = [
            ['Criterio', 'Score', 'Pts'],
            ['SROI (40%)', f"{self.resultado.score_sroi:.0f}", f"{self.resultado.contribucion_sroi:.1f}"],
            ['Stakeholders (25%)', f"{self.resultado.score_stakeholders:.0f}", f"{self.resultado.contribucion_stakeholders:.1f}"],
            ['Probabilidad (20%)', f"{self.resultado.score_probabilidad:.0f}", f"{self.resultado.contribucion_probabilidad:.1f}"],
            ['Riesgos (15%)', f"{self.resultado.score_riesgos:.0f}", f"{self.resultado.contribucion_riesgos:.1f}"]
        ]
        
        criterios_table = Table(criterios_compactos, colWidths=[3*inch, 1.5*inch, 1.5*inch])
        criterios_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        
        story.append(criterios_table)
        story.append(Spacer(1, 0.1*inch))
        
        # Conclusión
        if self.resultado.score_total >= 85:
            conclusion = "✓ PROYECTO DE MUY ALTA PRIORIDAD - Recomendado para aprobación inmediata"
        elif self.resultado.score_total >= 70:
            conclusion = "✓ PROYECTO DE ALTA PRIORIDAD - Recomendado para aprobación"
        elif self.resultado.score_total >= 50:
            conclusion = "⚠ PROYECTO DE PRIORIDAD MEDIA - Requiere evaluación detallada"
        else:
            conclusion = "✗ PROYECTO DE BAJA PRIORIDAD - No recomendado"
        
        conclusion_style = ParagraphStyle(
            'Conclusion',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceAfter=5
        )
        
        story.append(Paragraph(conclusion, conclusion_style))
        story.append(Spacer(1, 0.05*inch))
        
        # Footer
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        
        story.append(Paragraph(f"Evaluado: {self.fecha_evaluacion} | Sistema: Motor Arquitectura C v1.0", footer_style))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()
