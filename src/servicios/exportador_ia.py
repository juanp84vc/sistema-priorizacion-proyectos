"""
Servicio de exportación de análisis IA a diferentes formatos.
"""
import markdown
from datetime import datetime
from typing import Dict, List, Optional
from pathlib import Path
import io

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ExportadorIA:
    """Exporta análisis de IA a diferentes formatos (PDF, Markdown, DOCX)."""

    def __init__(self):
        """Inicializa el exportador."""
        self.output_dir = Path(__file__).parent.parent.parent / 'exports'
        self.output_dir.mkdir(exist_ok=True)

    def exportar_a_markdown(self, consulta: Dict, incluir_metadata: bool = True) -> str:
        """
        Exporta una consulta a formato Markdown.

        Args:
            consulta: Diccionario con los datos de la consulta
            incluir_metadata: Si incluir metadatos (fecha, LLM usado, etc.)

        Returns:
            String con el contenido en Markdown
        """
        md_content = []

        # Encabezado
        md_content.append(f"# {consulta.get('tipo_analisis', 'Análisis').replace('_', ' ').title()}")
        md_content.append("")

        # Metadata
        if incluir_metadata:
            md_content.append("## 📋 Información del Análisis")
            md_content.append("")

            if consulta.get('proyecto_nombre'):
                md_content.append(f"**Proyecto:** {consulta['proyecto_nombre']}")

            if consulta.get('timestamp'):
                fecha = datetime.fromisoformat(consulta['timestamp'])
                md_content.append(f"**Fecha:** {fecha.strftime('%d/%m/%Y %H:%M')}")

            if consulta.get('llm_provider'):
                llm_info = consulta['llm_provider']
                if consulta.get('llm_model'):
                    llm_info += f" ({consulta['llm_model']})"
                md_content.append(f"**LLM:** {llm_info}")

            md_content.append("")
            md_content.append("---")
            md_content.append("")

        # Pregunta
        md_content.append("## 🤔 Pregunta")
        md_content.append("")
        md_content.append(consulta.get('pregunta', ''))
        md_content.append("")

        # Respuesta
        md_content.append("## 💡 Análisis")
        md_content.append("")
        md_content.append(consulta.get('respuesta', ''))
        md_content.append("")

        # Footer
        if incluir_metadata:
            md_content.append("---")
            md_content.append("")
            md_content.append(f"*Generado por Sistema de Priorización de Proyectos - {datetime.now().strftime('%d/%m/%Y')}*")

        return '\n'.join(md_content)

    def guardar_markdown(self, consulta: Dict, nombre_archivo: Optional[str] = None) -> str:
        """
        Guarda una consulta como archivo Markdown.

        Args:
            consulta: Diccionario con los datos de la consulta
            nombre_archivo: Nombre del archivo (opcional, se genera automáticamente si no se proporciona)

        Returns:
            Ruta completa del archivo guardado
        """
        if nombre_archivo is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            tipo = consulta.get('tipo_analisis', 'analisis')
            nombre_archivo = f"{tipo}_{timestamp}.md"

        # Asegurar que tiene extensión .md
        if not nombre_archivo.endswith('.md'):
            nombre_archivo += '.md'

        ruta_archivo = self.output_dir / nombre_archivo
        contenido = self.exportar_a_markdown(consulta)

        with open(ruta_archivo, 'w', encoding='utf-8') as f:
            f.write(contenido)

        return str(ruta_archivo)

    def exportar_a_pdf(self, consulta: Dict, incluir_metadata: bool = True) -> bytes:
        """
        Exporta una consulta a formato PDF.

        Args:
            consulta: Diccionario con los datos de la consulta
            incluir_metadata: Si incluir metadatos

        Returns:
            Bytes del PDF generado

        Raises:
            ImportError: Si reportlab no está instalado
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab no está instalado. Ejecuta: pip install reportlab")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)

        # Estilos
        styles = getSampleStyleSheet()
        story = []

        # Estilo personalizado para el título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2E4057'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        # Estilo para secciones
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#3498db'),
            spaceAfter=12,
            spaceBefore=12
        )

        # Estilo para texto normal
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )

        # Título
        tipo_analisis = consulta.get('tipo_analisis', 'Análisis').replace('_', ' ').title()
        story.append(Paragraph(tipo_analisis, title_style))
        story.append(Spacer(1, 0.2*inch))

        # Metadata
        if incluir_metadata:
            metadata_data = []

            if consulta.get('proyecto_nombre'):
                metadata_data.append(['Proyecto:', consulta['proyecto_nombre']])

            if consulta.get('timestamp'):
                fecha = datetime.fromisoformat(consulta['timestamp'])
                metadata_data.append(['Fecha:', fecha.strftime('%d/%m/%Y %H:%M')])

            if consulta.get('llm_provider'):
                llm_info = consulta['llm_provider']
                if consulta.get('llm_model'):
                    llm_info += f" ({consulta['llm_model']})"
                metadata_data.append(['LLM:', llm_info])

            if metadata_data:
                metadata_table = Table(metadata_data, colWidths=[1.5*inch, 4.5*inch])
                metadata_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#34495e')),
                    ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(metadata_table)
                story.append(Spacer(1, 0.3*inch))

        # Pregunta
        story.append(Paragraph("Pregunta", heading_style))
        pregunta_texto = consulta.get('pregunta', '').replace('\n', '<br/>')
        story.append(Paragraph(pregunta_texto, normal_style))
        story.append(Spacer(1, 0.2*inch))

        # Respuesta
        story.append(Paragraph("Análisis", heading_style))

        # Convertir Markdown a texto plano mejorado para PDF
        respuesta = consulta.get('respuesta', '')

        # Procesar la respuesta línea por línea
        for linea in respuesta.split('\n'):
            linea = linea.strip()

            if not linea:
                story.append(Spacer(1, 0.1*inch))
                continue

            # Encabezados Markdown
            if linea.startswith('###'):
                texto = linea.replace('###', '').strip()
                story.append(Paragraph(texto, styles['Heading3']))
            elif linea.startswith('##'):
                texto = linea.replace('##', '').strip()
                story.append(Paragraph(texto, styles['Heading2']))
            elif linea.startswith('#'):
                texto = linea.replace('#', '').strip()
                story.append(Paragraph(texto, styles['Heading1']))
            # Listas
            elif linea.startswith('- ') or linea.startswith('* '):
                texto = '• ' + linea[2:]
                story.append(Paragraph(texto, normal_style))
            elif linea.startswith(tuple(f'{i}.' for i in range(10))):
                story.append(Paragraph(linea, normal_style))
            # Texto normal
            else:
                # Reemplazar negritas (**texto** -> <b>texto</b>)
                linea = linea.replace('**', '<b>', 1)
                if '<b>' in linea:
                    linea = linea.replace('**', '</b>', 1)

                story.append(Paragraph(linea, normal_style))

        # Footer
        story.append(Spacer(1, 0.3*inch))
        footer_text = f"Generado por Sistema de Priorización de Proyectos - {datetime.now().strftime('%d/%m/%Y')}"
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph(footer_text, footer_style))

        # Construir PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        return pdf_bytes

    def guardar_pdf(self, consulta: Dict, nombre_archivo: Optional[str] = None) -> str:
        """
        Guarda una consulta como archivo PDF.

        Args:
            consulta: Diccionario con los datos de la consulta
            nombre_archivo: Nombre del archivo (opcional)

        Returns:
            Ruta completa del archivo guardado
        """
        if nombre_archivo is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            tipo = consulta.get('tipo_analisis', 'analisis')
            nombre_archivo = f"{tipo}_{timestamp}.pdf"

        # Asegurar que tiene extensión .pdf
        if not nombre_archivo.endswith('.pdf'):
            nombre_archivo += '.pdf'

        ruta_archivo = self.output_dir / nombre_archivo
        pdf_bytes = self.exportar_a_pdf(consulta)

        with open(ruta_archivo, 'wb') as f:
            f.write(pdf_bytes)

        return str(ruta_archivo)

    def exportar_multiples_a_pdf(self, consultas: List[Dict],
                                 titulo: str = "Reporte de Análisis") -> bytes:
        """
        Exporta múltiples consultas a un solo PDF.

        Args:
            consultas: Lista de consultas a exportar
            titulo: Título del reporte

        Returns:
            Bytes del PDF generado
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab no está instalado. Ejecuta: pip install reportlab")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)

        styles = getSampleStyleSheet()
        story = []

        # Portada
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2E4057'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(titulo, title_style))
        story.append(Spacer(1, 0.5*inch))

        subtitle = f"Generado el {datetime.now().strftime('%d de %B de %Y')}"
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph(subtitle, subtitle_style))
        story.append(Paragraph(f"Total de análisis: {len(consultas)}", subtitle_style))
        story.append(PageBreak())

        # Agregar cada consulta
        for i, consulta in enumerate(consultas, 1):
            # Separador entre consultas
            if i > 1:
                story.append(PageBreak())

            # Número de análisis
            num_style = ParagraphStyle(
                'AnalysisNumber',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey
            )
            story.append(Paragraph(f"Análisis {i} de {len(consultas)}", num_style))
            story.append(Spacer(1, 0.2*inch))

            # Usar el método existente pero sin metadata completa
            # (para evitar repetición en reportes grandes)
            consulta_modificada = consulta.copy()
            pdf_individual = self.exportar_a_pdf(consulta_modificada, incluir_metadata=True)

        # Construir PDF final
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        return pdf_bytes

    def exportar_a_docx(self, consulta: Dict, incluir_metadata: bool = True) -> bytes:
        """
        Exporta una consulta a formato Word (DOCX).

        Args:
            consulta: Diccionario con los datos de la consulta
            incluir_metadata: Si incluir metadatos

        Returns:
            Bytes del documento Word generado

        Raises:
            ImportError: Si python-docx no está instalado
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecuta: pip install python-docx")

        doc = Document()

        # Título
        tipo_analisis = consulta.get('tipo_analisis', 'Análisis').replace('_', ' ').title()
        titulo = doc.add_heading(tipo_analisis, 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if incluir_metadata:
            doc.add_paragraph()  # Espacio

            if consulta.get('proyecto_nombre'):
                p = doc.add_paragraph()
                p.add_run('Proyecto: ').bold = True
                p.add_run(consulta['proyecto_nombre'])

            if consulta.get('timestamp'):
                fecha = datetime.fromisoformat(consulta['timestamp'])
                p = doc.add_paragraph()
                p.add_run('Fecha: ').bold = True
                p.add_run(fecha.strftime('%d/%m/%Y %H:%M'))

            if consulta.get('llm_provider'):
                llm_info = consulta['llm_provider']
                if consulta.get('llm_model'):
                    llm_info += f" ({consulta['llm_model']})"
                p = doc.add_paragraph()
                p.add_run('LLM: ').bold = True
                p.add_run(llm_info)

            doc.add_paragraph('_' * 50)  # Línea separadora

        # Pregunta
        doc.add_heading('Pregunta', level=1)
        doc.add_paragraph(consulta.get('pregunta', ''))

        # Respuesta
        doc.add_heading('Análisis', level=1)

        # Procesar respuesta (markdown básico)
        respuesta = consulta.get('respuesta', '')
        for linea in respuesta.split('\n'):
            linea = linea.strip()

            if not linea:
                doc.add_paragraph()
                continue

            # Encabezados
            if linea.startswith('###'):
                texto = linea.replace('###', '').strip()
                doc.add_heading(texto, level=3)
            elif linea.startswith('##'):
                texto = linea.replace('##', '').strip()
                doc.add_heading(texto, level=2)
            elif linea.startswith('#'):
                texto = linea.replace('#', '').strip()
                doc.add_heading(texto, level=1)
            # Listas
            elif linea.startswith('- ') or linea.startswith('* '):
                texto = linea[2:]
                doc.add_paragraph(texto, style='List Bullet')
            elif linea.startswith(tuple(f'{i}.' for i in range(10))):
                doc.add_paragraph(linea, style='List Number')
            # Texto normal
            else:
                p = doc.add_paragraph()
                # Procesar negritas básicas
                partes = linea.split('**')
                for i, parte in enumerate(partes):
                    if i % 2 == 0:
                        p.add_run(parte)
                    else:
                        p.add_run(parte).bold = True

        # Footer
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        footer = doc.add_paragraph(
            f'Generado por Sistema de Priorización de Proyectos - {datetime.now().strftime("%d/%m/%Y")}'
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Guardar en buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        return docx_bytes

    def guardar_docx(self, consulta: Dict, nombre_archivo: Optional[str] = None) -> str:
        """
        Guarda una consulta como archivo Word.

        Args:
            consulta: Diccionario con los datos de la consulta
            nombre_archivo: Nombre del archivo (opcional)

        Returns:
            Ruta completa del archivo guardado
        """
        if nombre_archivo is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            tipo = consulta.get('tipo_analisis', 'analisis')
            nombre_archivo = f"{tipo}_{timestamp}.docx"

        # Asegurar que tiene extensión .docx
        if not nombre_archivo.endswith('.docx'):
            nombre_archivo += '.docx'

        ruta_archivo = self.output_dir / nombre_archivo
        docx_bytes = self.exportar_a_docx(consulta)

        with open(ruta_archivo, 'wb') as f:
            f.write(docx_bytes)

        return str(ruta_archivo)
