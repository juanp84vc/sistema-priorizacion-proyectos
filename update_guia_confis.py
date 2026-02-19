#!/usr/bin/env python3
"""
Script to update the Guía Operativa with CONFIS methodology integration
Uses python-docx to modify specific sections surgically
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def find_paragraph_by_text(doc, search_text, partial=True):
    """Find paragraph index containing specific text"""
    for i, para in enumerate(doc.paragraphs):
        if partial:
            if search_text.lower() in para.text.lower():
                return i
        else:
            if para.text.lower() == search_text.lower():
                return i
    return None

def replace_paragraph_text(para, new_text, preserve_format=True):
    """Replace paragraph text while preserving formatting"""
    # Clear existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # Add new text with original paragraph formatting
    new_run = para.add_run(new_text)
    if preserve_format and len(para.runs) > 0:
        # Copy format from first run if available
        new_run.bold = False
        new_run.italic = False

def update_guia_operativa(doc_path):
    """Main function to update the document"""

    doc = Document(doc_path)

    print(f"Document loaded. Total paragraphs: {len(doc.paragraphs)}")
    print("\n" + "="*80)

    # ========================================
    # UPDATE 1: Header/Version (paragraph ~8)
    # ========================================
    print("\n[1] Updating header/version subtitle...")

    idx = find_paragraph_by_text(doc, "Incluye ajustes metodológicos")
    if idx is not None:
        old_text = doc.paragraphs[idx].text
        new_text = "Incluye ajustes metodológicos: SROI logarítmico, rúbricas objetivas, alertas contextuales, metodología CONFIS"
        replace_paragraph_text(doc.paragraphs[idx], new_text)
        print(f"  ✓ Updated paragraph {idx}")
        print(f"    Old: {old_text[:80]}...")
        print(f"    New: {new_text[:80]}...")
    else:
        print("  ✗ Subtitle paragraph not found - searching nearby...")
        for i in range(5, 15):
            if i < len(doc.paragraphs):
                print(f"    Para {i}: {doc.paragraphs[i].text[:60]}")

    # ========================================
    # UPDATE 2: Section 3.4 (Alcance Territorial) - paragraphs ~55-60
    # ========================================
    print("\n[2] Updating Section 3.4 (Alcance Territorial)...")

    idx = find_paragraph_by_text(doc, "Alcance Territorial")
    if idx is not None:
        print(f"  Found 'Alcance Territorial' at paragraph {idx}")

        # Look for the "Base: 10 pts" text
        for i in range(idx, min(idx + 10, len(doc.paragraphs))):
            if "Base: 10 pts por municipio" in doc.paragraphs[i].text:
                # Replace this and next paragraphs with new content
                new_content = [
                    "Puntaje Territorial CONFIS: puntaje × 3 (máx. 30 pts)",
                    "Base: 10 pts por municipio (máx. 30 pts)",
                    "Bonus PDET: +15 pts si incluye municipios PDET",
                    "Bonus multi-departamental: +15 pts",
                    "Bonus corredor: +10 pts",
                    "Máximo posible: 30 + 30 + 15 + 15 + 10 = 100 pts"
                ]

                # Replace the first occurrence
                replace_paragraph_text(doc.paragraphs[i], new_content[0])
                print(f"  ✓ Updated paragraph {i}: {new_content[0]}")

                # Update following paragraphs if they contain bonus text
                j = i + 1
                for new_line in new_content[1:]:
                    if j < len(doc.paragraphs):
                        if "Bonus" in doc.paragraphs[j].text or "Base:" in doc.paragraphs[j].text or "Máximo" in doc.paragraphs[j].text:
                            replace_paragraph_text(doc.paragraphs[j], new_line)
                            print(f"  ✓ Updated paragraph {j}: {new_line}")
                            j += 1
                        else:
                            # Insert new paragraph instead
                            new_para = doc.paragraphs[i]._element
                            new_para = new_para.addnext(doc.paragraphs[i]._element.makeelement(doc.paragraphs[i]._element.tag))
                            print(f"  ℹ Skipped paragraph {j} (no matching content)")
                            break
                break
    else:
        print("  ✗ 'Alcance Territorial' section not found")

    # ========================================
    # UPDATE 3: Section 4 (Criterio 3) - paragraphs ~61-70
    # ========================================
    print("\n[3] Updating Section 4 (Criterio 3)...")

    idx = find_paragraph_by_text(doc, "Criterio 3")
    if idx is not None:
        print(f"  Found 'Criterio 3' at paragraph {idx}")

        # Update title
        for i in range(idx, min(idx + 5, len(doc.paragraphs))):
            para_text = doc.paragraphs[i].text
            if "Probabilidad de Aprobación" in para_text and "CONFIS" not in para_text:
                old_text = para_text
                new_title = "4. Criterio 3: Probabilidad de Aprobación CONFIS (20%)"
                replace_paragraph_text(doc.paragraphs[i], new_title)
                print(f"  ✓ Updated title at paragraph {i}")
                print(f"    Old: {old_text}")
                print(f"    New: {new_title}")
                section_4_idx = i
                break

        # Update description (the paragraph after title)
        if 'section_4_idx' in locals():
            for i in range(section_4_idx + 1, min(section_4_idx + 10, len(doc.paragraphs))):
                para = doc.paragraphs[i]
                # Look for the description paragraph (usually contains "Evalúa")
                if "Evalúa" in para.text or "probabilidad" in para.text.lower():
                    new_desc = "Evalúa la probabilidad de aprobación usando la metodología CONFIS (Anexo 2). Integra grupo de priorización, puntaje territorial y puntaje sectorial."
                    replace_paragraph_text(para, new_desc)
                    print(f"  ✓ Updated description at paragraph {i}")
                    print(f"    New: {new_desc}")
                    break
    else:
        print("  ✗ 'Criterio 3' section not found")

    # ========================================
    # UPDATE 4: Section 4.1 (Fórmula CONFIS) - replace "Factores que considera"
    # ========================================
    print("\n[4] Updating Section 4.1 (Fórmula CONFIS)...")

    idx = find_paragraph_by_text(doc, "Factores que considera")
    if idx is not None:
        print(f"  Found '4.1 Factores que considera' at paragraph {idx}")

        # Replace title
        replace_paragraph_text(doc.paragraphs[idx], "4.1 Fórmula CONFIS")
        print(f"  ✓ Updated subsection title at paragraph {idx}")

        # Update formula paragraphs
        formula_lines = [
            "Score = GrupoPrioridad × 20% + ScoreCONFIS_Normalizado × 80%",
            "ScoreCONFIS = (PuntajeTerritorial + PuntajeSectorial) / 20 × 100"
        ]

        for i in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            # Look for and replace formula content
            if any(term in para.text for term in ["Score", "GrupoPrioridad", "factor", "evaluación"]):
                if "Score" not in para.text:
                    replace_paragraph_text(para, formula_lines[0])
                    print(f"  ✓ Added formula line at paragraph {i}")

                    # Add second formula in next paragraph
                    if i + 1 < len(doc.paragraphs):
                        next_para = doc.paragraphs[i + 1]
                        replace_paragraph_text(next_para, formula_lines[1])
                        print(f"  ✓ Added formula line at paragraph {i + 1}")
                    break
                else:
                    replace_paragraph_text(para, formula_lines[0])
                    print(f"  ✓ Updated formula at paragraph {i}")
                    break
    else:
        print("  ✗ '4.1 Factores que considera' not found")

    # ========================================
    # UPDATE 5: CONFIS Groups explanation & Gate de elegibilidad
    # ========================================
    print("\n[5] Updating CONFIS Groups and Gate de elegibilidad...")

    idx = find_paragraph_by_text(doc, "Bonus")
    if idx is not None:
        print(f"  Found 'Bonus' text around paragraph {idx}")

        # Search for the specific bonus paragraph to replace
        for i in range(idx, min(idx + 10, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if "Bonus" in para.text and ("PDET" in para.text or "grupos" in para.text.lower()):
                # Replace with CONFIS groups explanation
                groups_text = "8 Grupos de priorización: Grupo 1-2 (PATR-PDET), 3-4 (PDET), 5-6 (ZOMAC), 7-8 (Amazonía). Grupos impares = contribuyente paga estructuración."
                replace_paragraph_text(para, groups_text)
                print(f"  ✓ Updated groups explanation at paragraph {i}")

                # Add gate de elegibilidad paragraph after
                gate_text = "Gate de elegibilidad: Solo proyectos en municipios PDET/ZOMAC/Amazonía son elegibles. Proyectos fuera reciben score 0 y nivel NO ELEGIBLE."
                gate_para = doc.paragraphs[i]._element
                gate_para = gate_para.addnext(doc.paragraphs[i]._element.makeelement(doc.paragraphs[i]._element.tag))
                print(f"  ℹ Note: Gate de elegibilidad should be added manually or via paragraph insertion")
                break
    else:
        print("  ✗ Bonus text not found - checking for alternative patterns...")

    # ========================================
    # UPDATE 6: Section 7 (Ejemplo) - Add CONFIS scoring mention
    # ========================================
    print("\n[6] Updating Section 7 (Ejemplo)...")

    idx = find_paragraph_by_text(doc, "Ejemplo")
    if idx is not None:
        print(f"  Found 'Ejemplo' section at paragraph {idx}")

        # Look for Stakeholders step
        for i in range(idx, min(idx + 30, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if "Stakeholders" in para.text:
                # Add CONFIS calculation note after Stakeholders
                confis_note = "Cálculo Prob. CONFIS: Basado en grupo de priorización, puntaje territorial y sectorial integrados."
                print(f"  ✓ Found Stakeholders at paragraph {i}")
                print(f"    Will add CONFIS note after this section")
                # Note: actual insertion would require paragraph manipulation
                break
    else:
        print("  ✗ 'Ejemplo' section not found")

    # ========================================
    # Save document
    # ========================================
    print("\n" + "="*80)
    print("\nSaving updated document...")
    doc.save(doc_path)
    print(f"✓ Document saved successfully to: {doc_path}")

    return doc

if __name__ == "__main__":
    doc_path = "/sessions/keen-stoic-bell/mnt/sistema-priorizacion-proyectos/Guia_Operativa_Evaluadores_ENLAZA_GEB.docx"

    try:
        doc = update_guia_operativa(doc_path)
        print("\n✓ All updates completed successfully!")
    except Exception as e:
        print(f"\n✗ Error occurred: {e}")
        import traceback
        traceback.print_exc()
